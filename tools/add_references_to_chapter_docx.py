from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_With_References.docx"


REFERENCES = [
    "Abao, E. F., Jr., Baja, C. V. U., & Taqueban, C. O. (2024). Motormania's web-based streamline point-of-sales inventory management with predictive analytics [Undergraduate manuscript]. University of Science and Technology of Southern Philippines.",
    "Accad, M. J. R., Bantayan, R. M., Calma, A. A., Jr., Maquinana, K. E., Tanael, D. V., & Maupay, A. C. M. (2023). Integrated inventory management and asset tracking system with user-centric computer kiosk interface. World Journal of Advanced Research and Reviews, 20(3), 270-276. https://doi.org/10.30574/wjarr.2023.20.3.2471",
    "Baimukhamedova, A. (2022). Digital transformation in the Western Balkans as an opportunity for managing innovation in small and medium businesses: Challenges and opportunities. IFAC-PapersOnLine, 55(39), 60-65. https://doi.org/10.1016/j.ifacol.2022.12.011",
    "Cadiz Lagmay, J. V. P., & Palaoag, T. D. (2024). ProCoMon: A web-based project procurement management plan (PPMP) consolidation and monitoring system for Nueva Vizcaya State University. Journal of Electrical Systems, 20(4s), 1234-1240. https://doi.org/10.52783/jes.2168",
    "Choudhury, T. M. (2022). Web-based POS and inventory management system for SMEs, retail stores, and restaurants [Undergraduate project, Independent University, Bangladesh]. IUB Academic Repository. https://ar.iub.edu.bd/handle/11348/787",
    "Divina, J., Olan, A., Perez, N. C., Acepcion, R., Sarmiento, M., & Cabututan, G. (2025). Data-driven point-of-sale and inventory system for Pastil Sa Tabi: Integrating sales forecasting algorithms and predictive analytics. International Journal of Research and Innovation in Applied Science, 10(10), 826-838. https://doi.org/10.51584/IJRIAS.2025.1010000066",
    "Elec, D. J. V., & Martin, M. L. A. (2025). Development of a barcode-based inventory system for AMT laboratory tools. Diversitas Journal, 10(Special Issue 1). https://doi.org/10.48017/dj.v10ispecial_1.3177",
    "Flores-Alhaddad, J. J. C., & Cabudol, E. G. (2024). Inventory management techniques of heat stores in Metro Manila: Basis for improvement of inventory management system. IOER International Multidisciplinary Research Journal, 6(3). https://www.ioer-imrj.com/inventory-management-techniques-of-heat-stores-in-metro-manila-basis-for-improvement-of-inventory-management-system/",
    "Gouveia, F. D., & Sao Mamede, H. (2022). Digital transformation for SMEs in the retail industry. Procedia Computer Science, 204, 671-681. https://doi.org/10.1016/j.procs.2022.08.081",
    "International Organization for Standardization. (2023). Systems and software engineering: Systems and software quality requirements and evaluation (SQuaRE): Product quality model (ISO/IEC 25010:2023). ISO.",
    "Karmarkar, B., Nimsatkar, S. S., Sadamwar, A., et al. (2024). Inventory management system. International Journal of Advanced Research in Science, Communication and Technology, 4(4), 632-636. https://doi.org/10.48175/ijarsct-18364",
    "Kovalyshyn, M., & Paramud, Y. (2024). Web application for warehouse management system. Computer Systems and Networks, 6(2), 107-122. https://science.lpnu.ua/csn/all-volumes-and-issues/volume-6-number-2-2024/web-application-warehouse-management-system",
    "Lapada, A. A. (2023). Development and evaluation of a web-based management information system for a planning unit of a state university in the Philippines. Journal of University of Shanghai for Science and Technology.",
    "Magallanes, C. A., Ortiz, M. N., Seville, M. N., Tejada, S. L. G., Tuliao, E. M., Eroy, N. G., & Buladaco, M. V. M. (2021). Analysis and design of a sales and inventory management information system for a motorcycle parts and accessories store. International Journal of Scientific Research and Engineering Development, 4(3).",
    "Melvin, M., Wiratama, J., Sutomo, R., & Sanjaya, S. A. (2023). A web-based point of sales for automotive component industry using rapid application development model. JOINS (Journal of Information System), 8(2), 167-176. https://doi.org/10.33633/joins.v8i2.9383",
    "Nguyen, Y. (2023). Warehouse item tracking and inventory control website [Bachelor's thesis, LAB University of Applied Sciences]. Theseus. https://www.theseus.fi/handle/10024/800714",
    "Ramadhani, I., Nindyasari, R., & Murti, A. C. (2025). Design and development of a web-based point of sale system for small-scale retail management. Bit-Tech, 8(1), 181-189. https://garuda.kemdiktisaintek.go.id/documents/detail/5352374",
    "Schwaber, K., & Sutherland, J. (1995). Scrum development process. In J. Sutherland, C. Casanave, J. Miller, P. Patel, & G. Hollowell (Eds.), Business object design and implementation: OOPSLA '95 workshop proceedings. Springer.",
    "Shan, W., Wan, W., Chen, A., Ren, L., Sun, J., Fang, M., & Zhan, Y. (2023). 3D warehousing: Enabling intelligent warehousing visualization based on Three.js. In Signal and Information Processing, Networking and Computers (Lecture Notes in Electrical Engineering, Vol. 996, pp. 63-71). Springer. https://doi.org/10.1007/978-981-19-9968-0_8",
    "Tanaman, M. T., Baylosis, J. L. A., Abiles, B. J. A., Catungal, M. L. P., & Encarnacion, P. C. (2023). Web-based inventory management system. International Journal of Science and Applied Information Technology, 12(5), 44-48. https://doi.org/10.30534/ijsait/2023/021252023",
    "Wahyudi, H., & Hendriawan, D. (2023). Web based sales transaction and report system design: Study at a motorcycle spare parts company. Electronic, Business, Management and Technology Journal, 1(2), 129-135. https://doi.org/10.55208/ebmtj.v1i2.125",
]


def format_run(run, *, bold: bool = False, size: int = 11):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Times New Roman")
    font.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(font)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    properties.append(size)

    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_reference_paragraph(doc: Document, reference: str):
    paragraph = doc.add_paragraph()
    url_match = re.search(r"https?://\S+$", reference)
    if url_match:
        before = reference[: url_match.start()]
        url = url_match.group(0)
        run = paragraph.add_run(before)
        format_run(run, size=11)
        add_hyperlink(paragraph, url, url)
    else:
        run = paragraph.add_run(reference)
        format_run(run, size=11)
    format_reference_paragraph(paragraph)


def format_reference_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        format_run(run, size=11)


def main():
    doc = Document(SOURCE_DOCX)

    if any(p.text.strip().upper().startswith("REFERENCES") for p in doc.paragraphs):
        raise SystemExit("The document already appears to contain a REFERENCES section.")

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("REFERENCES:")
    format_run(run, bold=True, size=12)
    heading.paragraph_format.space_after = Pt(12)

    for reference in REFERENCES:
        add_reference_paragraph(doc, reference)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
