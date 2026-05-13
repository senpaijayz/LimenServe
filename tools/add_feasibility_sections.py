from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Feasibility_Gantt.docx"


FEASIBILITY_SECTIONS = [
    (
        "Operational Feasibility",
        "The proposed LimenServe system is operationally feasible because it directly addresses the current manual process of Limen Auto Parts Center. At present, employees still depend on handwritten receipts, paper-based inventory logs, manual quotation preparation, and staff familiarity when locating products in the stockroom. These practices may cause delays in responding to customer inquiries, difficulty in checking available stocks, inconsistent quotation preparation, and delayed report consolidation. Through LimenServe, daily operations can become more organized because inventory records, quotations, sales transactions, service orders, reports, and stockroom location details are managed in one centralized web-based platform.",
        "The system also supports the existing workflow of the business because it does not require the shop to fully change its operational structure. Instead, it improves the current process by providing digital tools for product monitoring, barcode-assisted identification, automated cost estimation, point-of-sale transactions, and service order tracking. Authorized users such as administrators, cashiers, and stock personnel can perform their assigned tasks based on their roles. Therefore, the system is feasible in terms of actual use because it can reduce manual work, improve transaction speed, minimize record inconsistencies, and help the business provide faster and more accurate service to customers.",
    ),
    (
        "Technical Feasibility",
        "The proposed system is technically feasible because the required hardware and software resources are available and practical for a small retail and service-oriented business. The system can run through standard desktop or laptop computers using a web browser and an internet connection. Additional devices such as a barcode scanner and receipt printer may be used to support product identification and transaction recording. Since the system is web-based, users do not need to install a separate desktop application, making access and maintenance more manageable.",
        "The technologies identified for the development of LimenServe include React.js for the frontend interface, Node.js for backend processing, Supabase for database management and authentication, and Three.js for the 3D stockroom visualization module. These technologies are suitable for creating a responsive, database-driven, and role-based management information system. The system can also be deployed as a Progressive Web App (PWA), allowing it to be accessed through modern browsers. Based on the availability of the required tools, devices, and development technologies, the proposed system is technically feasible.",
    ),
    (
        "Economic Feasibility",
        "The proposed system is economically feasible because it is designed to reduce the time and effort spent on manual recordkeeping, quotation preparation, inventory checking, and report generation. Although the development and implementation of the system may require costs for devices, deployment, training, and maintenance, these costs are reasonable compared with the long-term benefits that the system can provide. By improving record accuracy and reducing repetitive manual work, the business can lessen possible errors in pricing, stock monitoring, and transaction recording.",
        "The system also uses accessible web technologies and cloud-based services, which helps reduce the need for expensive physical infrastructure. Existing computers and internet access may be used, while optional hardware such as barcode scanners and receipt printers can be added depending on business needs. The economic value of the proposed system is shown through faster customer service, better monitoring of inventory, easier preparation of sales and service reports, and improved decision-making for the owner. Therefore, the project is economically feasible because its expected operational benefits justify the required implementation resources.",
    ),
    (
        "Schedule Feasibility",
        "The proposed system is schedule-feasible because the development activities are divided into manageable phases that follow the research and system development timeline. The schedule begins with proposal preparation and requirement gathering, followed by proposal defense, system design, development, testing, implementation, deployment, manuscript review, and final oral defense. Dividing the work into phases allows the researchers to monitor progress and complete the system according to the planned timeline.",
        "The Gantt chart below presents the major activities and target schedule for the development of LimenServe. It shows that the project activities are arranged in sequence, while some activities such as manuscript preparation, system development, testing, and evaluation may overlap to support continuous improvement. Based on the planned timeline, the proposed system can be completed within the academic development period.",
    ),
]


GANTT_COLUMNS = [
    "Activity",
    "Oct-Dec\n2025",
    "Jan\n2026",
    "Feb-Mar\n2026",
    "Mar-Dec\n2026",
    "Oct-Nov\n2026",
    "Oct-Dec\n2026",
    "Dec\n2026",
    "Nov 2026-\nFeb 2027",
    "Jan\n2027",
]


GANTT_ROWS = [
    ("Req.", "X", "", "", "", "", "", "", "", ""),
    ("Defense", "", "X", "", "", "", "", "", "", ""),
    ("Design", "", "", "X", "", "", "", "", "", ""),
    ("Dev.", "", "", "", "X", "", "", "", "", ""),
    ("Impl.", "", "", "", "", "X", "", "", "", ""),
    ("Testing", "", "", "", "", "", "X", "", "", ""),
    ("Deploy", "", "", "", "", "", "", "X", "", ""),
    ("Review", "", "", "", "", "", "", "", "X", ""),
    ("Final", "", "", "", "", "", "", "", "", "X"),
]


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table, text: str = ""):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    paragraph = table._parent.add_paragraph()
    paragraph._p = new_p
    if text:
        paragraph.add_run(text)
    return paragraph


def format_para(paragraph, *, heading: bool = False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if heading else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    if not heading:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    else:
        paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = heading


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False, fill: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(7)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_gantt_table(after_paragraph):
    table = insert_table_after(after_paragraph, 1, len(GANTT_COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, heading in enumerate(GANTT_COLUMNS):
        set_cell_text(table.rows[0].cells[idx], heading, bold=True, fill="D9EAF7")

    for row in GANTT_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx == 0:
                set_cell_text(cells[idx], value)
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif value == "X":
                set_cell_text(cells[idx], "X", bold=True, fill="1F4E79")
            else:
                set_cell_text(cells[idx], "")

    widths = [1.1, 0.62, 0.5, 0.58, 0.58, 0.58, 0.58, 0.5, 0.7, 0.5]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
    return table


def main():
    doc = Document(SOURCE_DOCX)
    if any(p.text.strip() == "Operational Feasibility" for p in doc.paragraphs):
        raise SystemExit("Detailed feasibility sections already appear to exist.")

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "Table 3.2. Feasibility analysis":
            anchor = p
            break
    if anchor is None:
        raise SystemExit("Could not locate Feasibility Analysis table caption.")

    cursor = anchor
    last_schedule_para = None
    for title, para1, para2 in FEASIBILITY_SECTIONS:
        cursor = insert_paragraph_after(cursor, title)
        format_para(cursor, heading=True)
        cursor = insert_paragraph_after(cursor, para1)
        format_para(cursor)
        cursor = insert_paragraph_after(cursor, para2)
        format_para(cursor)
        if title == "Schedule Feasibility":
            last_schedule_para = cursor

    if last_schedule_para is None:
        raise SystemExit("Schedule feasibility paragraph was not created.")

    table = add_gantt_table(last_schedule_para)
    caption = insert_paragraph_after_table(table, "Gantt Chart. LimenServe development schedule")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.italic = True

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
