from __future__ import annotations

import argparse
import html
import json
from pathlib import Path
from textwrap import wrap


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docs" / "wireframes" / "limenserve-system-wireframes"
FIG_DIR = OUT_DIR / "figures"
WIDTH = 1200
HEIGHT = 720


FIGURES = [
    {
        "slug": "public-home",
        "title": "Public Home Page Wireframe",
        "layout": "public_home",
        "description": (
            "Figure {n} shows the public landing page of LimenServe. It presents the shop identity, "
            "main navigation, search entry point, featured vehicle categories, product highlights, "
            "trust indicators, and call-to-action areas that guide customers to browse parts or request an estimate."
        ),
    },
    {
        "slug": "public-catalog",
        "title": "Public Parts Catalog Wireframe",
        "layout": "public_catalog",
        "description": (
            "Figure {n} shows the public parts catalog layout. It includes vehicle-first filters, search controls, "
            "product cards, matched service package suggestions, and a product detail area for customer browsing."
        ),
    },
    {
        "slug": "public-estimate-choice",
        "title": "Public Estimate Entry Wireframe",
        "layout": "estimate_choice",
        "description": (
            "Figure {n} shows the first phase of the public estimate workflow. Customers choose whether to build "
            "a new estimate or retrieve an existing quotation, then provide customer and vehicle details before "
            "continuing to parts and services."
        ),
    },
    {
        "slug": "public-estimate-builder",
        "title": "Public Estimate Builder Wireframe",
        "layout": "estimate_builder",
        "description": (
            "Figure {n} shows the public estimate builder. The layout combines part search, service selection, "
            "smart bundle recommendations, and a quotation cart so customers can review selected parts and services "
            "before finalizing the estimate."
        ),
    },
    {
        "slug": "public-quote-summary",
        "title": "Public Quotation Summary Wireframe",
        "layout": "quote_summary",
        "description": (
            "Figure {n} shows the public quotation summary and print preview structure. It presents customer and "
            "vehicle information, quotation line items, service lines, subtotals, estimated total, and printing controls."
        ),
    },
    {
        "slug": "public-service-orders",
        "title": "Public Service Orders Wireframe",
        "layout": "public_service_orders",
        "description": (
            "Figure {n} shows the public service-order information page. It explains the service workflow, from "
            "initial assessment to quotation, service handling, status tracking, and completion."
        ),
    },
    {
        "slug": "public-about",
        "title": "Public About Page Wireframe",
        "layout": "public_about",
        "description": (
            "Figure {n} shows the public about page of Limen Auto Parts Center. It contains the business story, "
            "operational overview, shop statistics, mechanics section, service values, and location information."
        ),
    },
    {
        "slug": "login",
        "title": "Staff Login Page Wireframe",
        "layout": "login",
        "description": (
            "Figure {n} shows the staff login page. It provides controlled access to the internal workspace through "
            "email and password fields, remember-me option, sign-in action, and administrative contact notice."
        ),
    },
    {
        "slug": "workspace-shell",
        "title": "Internal Workspace Shell Wireframe",
        "layout": "workspace_shell",
        "description": (
            "Figure {n} shows the common internal workspace shell used after authentication. It includes the sidebar "
            "navigation, top header, global search area, notification area, user identity section, and protected content area."
        ),
    },
    {
        "slug": "dashboard",
        "title": "Admin Dashboard Wireframe",
        "layout": "dashboard",
        "description": (
            "Figure {n} shows the dashboard module. It summarizes operational indicators through KPI cards, forecast "
            "widgets, sales trend charts, item highlights, recent transactions, and low-stock alerts."
        ),
    },
    {
        "slug": "inventory",
        "title": "Inventory Management Wireframe",
        "layout": "inventory",
        "description": (
            "Figure {n} shows the inventory management module. It provides search and filter controls, product listings, "
            "stock status indicators, barcode and label support, price-list import, and add-stock actions for inventory updates."
        ),
    },
    {
        "slug": "pos",
        "title": "Point of Sale Wireframe",
        "layout": "pos",
        "description": (
            "Figure {n} shows the point-of-sale module. It supports barcode or product search, cart building, quantity "
            "updates, payment entry, transaction total computation, and receipt preview for completed sales."
        ),
    },
    {
        "slug": "quotation",
        "title": "Internal Quotation Builder Wireframe",
        "layout": "quotation",
        "description": (
            "Figure {n} shows the internal quotation builder used by staff. It combines customer details, vehicle information, "
            "part selection, service selection, quotation line items, total computation, and print or save actions."
        ),
    },
    {
        "slug": "service-orders-admin",
        "title": "Service Orders Management Wireframe",
        "layout": "services_admin",
        "description": (
            "Figure {n} shows the internal service-order management module. It organizes service requests by status, "
            "supports search and filtering, displays order details, and allows staff to assign personnel or update progress."
        ),
    },
    {
        "slug": "stockroom",
        "title": "3D Stockroom Locator Wireframe",
        "layout": "stockroom",
        "description": (
            "Figure {n} shows the stockroom visualization module. It includes a 3D stockroom area, floor controls, "
            "product search, item location details, and route steps that help staff locate inventory inside the shop."
        ),
    },
    {
        "slug": "parts-mapping",
        "title": "Parts Mapping Admin Wireframe",
        "layout": "parts_mapping",
        "description": (
            "Figure {n} shows the parts mapping administration interface. It provides tools for editing the stockroom layout, "
            "mapping shelves and slots, modifying object properties, and saving layout changes used by the locator module."
        ),
    },
    {
        "slug": "reports",
        "title": "Reports and Analytics Wireframe",
        "layout": "reports",
        "description": (
            "Figure {n} shows the reporting and analytics module. It contains report tabs, date filters, summary cards, "
            "sales and inventory charts, tabular report results, and tools for encoding or editing historical sales data."
        ),
    },
    {
        "slug": "users",
        "title": "User Management Wireframe",
        "layout": "users",
        "description": (
            "Figure {n} shows the user management module. It displays user summary cards, search controls, role-based user "
            "records, and an add or edit user modal for maintaining authorized staff accounts."
        ),
    },
    {
        "slug": "cms",
        "title": "Content CMS Wireframe",
        "layout": "cms",
        "description": (
            "Figure {n} shows the content management module. It allows administrators to manage public pages, edit page sections, "
            "preview content, and publish updates that appear on the public LimenServe website."
        ),
    },
]


def esc(value: object) -> str:
    return html.escape(str(value), quote=True)


class Svg:
    def __init__(self, title: str):
        self.parts: list[str] = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{WIDTH}" height="{HEIGHT}" viewBox="0 0 {WIDTH} {HEIGHT}">',
            "<style>",
            "text{font-family:Arial,Helvetica,sans-serif;fill:#111;letter-spacing:0}",
            ".title{font-size:22px;font-weight:700}",
            ".label{font-size:14px;font-weight:700}",
            ".small{font-size:12px}",
            ".tiny{font-size:10px}",
            ".darktext{fill:#fff}",
            ".box{fill:#fff;stroke:#111;stroke-width:1.5}",
            ".soft{fill:#f7f7f7;stroke:#111;stroke-width:1.2}",
            ".dark{fill:#111;stroke:#111;stroke-width:1.2}",
            ".line{stroke:#111;stroke-width:1.4;fill:none}",
            ".thin{stroke:#111;stroke-width:1;fill:none}",
            "</style>",
            '<rect x="0" y="0" width="1200" height="720" fill="#fff"/>',
            f'<text x="600" y="34" text-anchor="middle" class="title">{esc(title)}</text>',
        ]

    def rect(self, x, y, w, h, label: str | None = None, cls="box", rx=0, anchor="middle", size="label"):
        self.parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" class="{cls}"/>')
        if label:
            text_cls = f"{size} darktext" if cls == "dark" else size
            self.multiline(label, x + w / 2 if anchor == "middle" else x + 12, y + h / 2, w - 18, anchor=anchor, cls=text_cls)

    def text(self, x, y, value, cls="small", anchor="start"):
        self.parts.append(f'<text x="{x}" y="{y}" text-anchor="{anchor}" class="{cls}">{esc(value)}</text>')

    def multiline(self, value, x, y, max_width, anchor="middle", cls="small", line_height=15):
        lines = []
        for paragraph in str(value).split("\n"):
            lines.extend(wrap(paragraph, width=max(8, int(max_width / 7.0))) or [""])
        start = y - ((len(lines) - 1) * line_height / 2)
        for i, line in enumerate(lines):
            self.text(x, start + i * line_height, line, cls=cls, anchor=anchor)

    def line(self, x1, y1, x2, y2, cls="line"):
        self.parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" class="{cls}"/>')

    def pill(self, x, y, w, h, label):
        self.rect(x, y, w, h, None, cls="soft", rx=12)
        self.text(x + w / 2, y + h / 2 + 4, label, cls="tiny", anchor="middle")

    def browser_frame(self):
        self.rect(40, 56, 1120, 620, None, cls="box", rx=0)
        self.rect(40, 56, 1120, 34, None, cls="soft", rx=0)
        for i, x in enumerate([60, 80, 100]):
            self.parts.append(f'<circle cx="{x}" cy="73" r="5" fill="#fff" stroke="#111" stroke-width="1"/>')
        self.rect(130, 64, 720, 18, "https://limen-serve.vercel.app", cls="box", rx=2, size="tiny")

    def finish(self):
        self.parts.append("</svg>")
        return "\n".join(self.parts)


def section_title(svg: Svg, x, y, w, label):
    svg.rect(x, y, w, 26, None, cls="dark")
    svg.parts.append(f'<text x="{x + w / 2}" y="{y + 18}" text-anchor="middle" font-size="12" font-weight="700" fill="#fff">{esc(label)}</text>')


def labeled_box(svg: Svg, x, y, w, h, label, detail=None):
    svg.rect(x, y, w, h, None, cls="box")
    svg.text(x + 10, y + 20, label, cls="label")
    if detail:
        svg.multiline(detail, x + 10, y + 43, w - 20, anchor="start", cls="small", line_height=15)


def public_header(svg: Svg):
    svg.browser_frame()
    svg.rect(60, 102, 1080, 24, "CONTACT BAR: phone numbers, store hours, location note", cls="soft", size="tiny")
    svg.rect(60, 132, 1080, 58, None, cls="box")
    svg.text(86, 166, "LIMEN / GENUINE AUTO PARTS", cls="label")
    nav = ["Home", "About", "Genuine Parts", "Get Estimate", "Service Orders", "Staff Portal"]
    x = 430
    for item in nav:
        svg.rect(x, 148, 92 if item != "Genuine Parts" else 118, 24, item, cls="soft", rx=2, size="tiny")
        x += 102 if item != "Genuine Parts" else 128


def app_shell(svg: Svg, active="Dashboard"):
    svg.browser_frame()
    svg.rect(60, 104, 190, 540, None, cls="soft")
    svg.text(88, 136, "LimenServe", cls="label")
    svg.text(88, 153, "Auto Parts MIS", cls="tiny")
    nav = ["Dashboard", "Point of Sale", "Inventory", "Quotation", "Service Orders", "3D Stockroom", "Reports", "User Management", "Content CMS"]
    y = 182
    for item in nav:
        cls = "box" if item == active else "soft"
        svg.rect(78, y, 144, 26, item, cls=cls, rx=2, size="tiny")
        y += 35
    svg.rect(78, 585, 144, 42, "USER ROLE\nLogout", cls="box", size="tiny")
    svg.rect(270, 104, 850, 52, None, cls="box")
    svg.text(292, 136, active, cls="label")
    svg.rect(750, 118, 180, 24, "GLOBAL SEARCH", cls="soft", size="tiny")
    svg.rect(942, 118, 70, 24, "ALERTS", cls="soft", size="tiny")
    svg.rect(1024, 118, 76, 24, "PROFILE", cls="soft", size="tiny")
    svg.rect(270, 174, 850, 450, None, cls="box")


def draw_public_home(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 500, 170, "HERO MESSAGE", "Genuine and aftermarket auto parts customers can trust.\nSearch field + Search Parts button\nVehicle quick chips")
    svg.rect(600, 212, 500, 170, "VEHICLE / STOREFRONT IMAGE PLACEHOLDER", cls="soft")
    for i, label in enumerate(["Search by Vehicle", "Featured Categories", "Fast Quotation"]):
        labeled_box(svg, 78 + i * 342, 402, 310, 72, label, "Short supporting text")
    for i, label in enumerate(["Mitsubishi Parts", "Service Packages", "Customer Trust", "Quote Request"]):
        labeled_box(svg, 78 + i * 255, 494, 230, 70, label, "Card content")
    svg.rect(78, 588, 1022, 46, "FOOTER / CONTACT / QUICK LINKS", cls="soft", size="small")


def draw_public_catalog(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 1022, 70, "PARTS CATALOG HEADING", "Search by part name, part number, or vehicle details")
    labeled_box(svg, 78, 298, 1022, 66, "VEHICLE-FIRST FILTERS", "Model dropdown | Year dropdown | Variant / engine | Reset filters")
    for i in range(3):
        for j in range(2):
            labeled_box(svg, 78 + i * 245, 386 + j * 86, 220, 64, f"PRODUCT CARD {i + 1 + j * 3}", "Part name\nPrice / stock / details")
    labeled_box(svg, 842, 386, 258, 150, "PRODUCT DETAIL PANEL", "Selected item\nCompatibility\nAdd to estimate")
    svg.rect(842, 554, 258, 70, "SMART BUNDLE SUGGESTIONS", cls="soft", size="small")


def draw_estimate_choice(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 1022, 66, "QUOTATION ENTRY", "Choose whether to build a new estimate or retrieve a saved quote")
    labeled_box(svg, 78, 302, 300, 92, "BUILD ESTIMATE CARD", "Pick parts, add services, and build a draft quotation")
    labeled_box(svg, 408, 302, 300, 92, "RETRIEVE QUOTE CARD", "Find a quotation number and open printable preview")
    labeled_box(svg, 738, 302, 362, 92, "PHASE STEPPER", "Customer + Vehicle > Parts and Services > Summary")
    labeled_box(svg, 78, 424, 500, 136, "CUSTOMER DETAILS FORM", "Full name\nPhone number\nEmail address")
    labeled_box(svg, 600, 424, 500, 136, "VEHICLE DETAILS FORM", "Model\nYear\nPlate number\nService concern")
    svg.rect(78, 586, 1022, 40, "CONTINUE TO PARTS AND SERVICES", cls="soft", size="label")


def draw_estimate_builder(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 680, 66, "PARTS AND SERVICES WORKSPACE", "Search parts, sort results, and add items to the quotation")
    labeled_box(svg, 780, 212, 320, 66, "QUOTE CART SUMMARY", "Selected parts + selected services")
    labeled_box(svg, 78, 300, 325, 210, "GENUINE PARTS LIST", "Search field\nFilter chips\nProduct cards with quantity controls")
    labeled_box(svg, 430, 300, 328, 210, "SERVICE CATALOG", "Service cards\nLabor cost\nSelect / remove")
    labeled_box(svg, 780, 300, 320, 210, "CART PANEL", "Line items\nServices subtotal\nEstimated total")
    svg.rect(78, 532, 680, 78, "GOOD / BETTER / BEST SMART BUNDLE RECOMMENDATIONS", cls="soft", size="small")
    svg.rect(780, 532, 320, 78, "BACK / CONTINUE TO SUMMARY", cls="soft", size="small")


def draw_quote_summary(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 500, 86, "CUSTOMER AND VEHICLE SUMMARY", "Customer, contact number, vehicle model, year, plate number")
    labeled_box(svg, 600, 212, 500, 86, "QUOTATION CONTROLS", "Save quotation | Print preview | Reset")
    labeled_box(svg, 78, 326, 700, 210, "QUOTATION LINE ITEMS TABLE", "Item / service name | Qty | Unit price | Line total")
    labeled_box(svg, 810, 326, 290, 210, "TOTALS PANEL", "Parts subtotal\nServices subtotal\nEstimated total\nQuotation number")
    svg.rect(78, 560, 1022, 62, "PRINT PREVIEW MODAL / FORMAL QUOTATION OUTPUT", cls="soft", size="small")


def draw_public_service_orders(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 1022, 76, "SERVICE ORDER WORKFLOW HERO", "Explains repair and installation handling for customers")
    for i, label in enumerate(["Initial Assessment", "Quotation", "Service Handling", "Completion"]):
        labeled_box(svg, 78 + i * 255, 318, 230, 112, f"STEP {i + 1}: {label}", "Customer-facing process details")
    labeled_box(svg, 78, 464, 500, 92, "STATUS TRACKING INFORMATION", "How customers can follow service progress")
    labeled_box(svg, 600, 464, 500, 92, "CONTACT / SERVICE CALL TO ACTION", "Visit store or coordinate with staff")
    svg.rect(78, 588, 1022, 36, "FOOTER", cls="soft", size="tiny")


def draw_public_about(svg: Svg):
    public_header(svg)
    labeled_box(svg, 78, 212, 500, 128, "ABOUT HERO", "Limen Auto Parts Center\nBusiness identity and operating context")
    labeled_box(svg, 600, 212, 500, 128, "SHOP / LOCATION VISUAL", "Store image or map placeholder")
    for i, label in enumerate(["13 Years", "Pasay City", "Mitsubishi Parts"]):
        labeled_box(svg, 78 + i * 342, 366, 310, 62, label, "Statistic card")
    labeled_box(svg, 78, 456, 500, 96, "OUR STORY AND OPERATIONS", "Sales area, stockroom, quotations, and daily process")
    labeled_box(svg, 600, 456, 500, 96, "MEET OUR MECHANICS", "Mechanic cards and role details")
    svg.rect(78, 582, 1022, 42, "MAP / CONTACT DETAILS", cls="soft", size="small")


def draw_login(svg: Svg):
    svg.browser_frame()
    svg.rect(80, 112, 180, 28, "BACK TO HOME", cls="soft", size="tiny")
    svg.rect(388, 130, 424, 430, None, cls="box")
    svg.text(600, 176, "Welcome Back", cls="title", anchor="middle")
    svg.text(600, 200, "Sign in to LimenServe MIS", cls="small", anchor="middle")
    svg.rect(438, 240, 324, 44, "EMAIL ADDRESS", cls="soft", size="tiny")
    svg.rect(438, 304, 324, 44, "PASSWORD", cls="soft", size="tiny")
    svg.rect(438, 368, 140, 24, "REMEMBER ME", cls="soft", size="tiny")
    svg.rect(622, 368, 140, 24, "FORGOT PASSWORD", cls="soft", size="tiny")
    svg.rect(438, 418, 324, 42, "SIGN IN", cls="dark", size="label")
    svg.rect(438, 486, 324, 44, "CONTACT ADMIN NOTICE", cls="soft", size="small")
    svg.text(600, 602, "(c) 2026 Limen Auto Parts Center", cls="tiny", anchor="middle")


def draw_workspace_shell(svg: Svg):
    app_shell(svg, "Protected Workspace")
    labeled_box(svg, 292, 196, 260, 110, "SIDEBAR NAVIGATION", "Role-based menu\nCollapsible desktop\nMobile overlay")
    labeled_box(svg, 578, 196, 250, 110, "TOP HEADER", "Current page title\nGlobal search\nNotifications")
    labeled_box(svg, 852, 196, 226, 110, "USER CONTEXT", "Profile initials\nRole display\nLogout")
    labeled_box(svg, 292, 334, 786, 226, "MODULE CONTENT AREA", "All internal modules render here after authentication.\nUnauthorized users are redirected to the login page or restricted message.")


def draw_dashboard(svg: Svg):
    app_shell(svg, "Dashboard")
    labeled_box(svg, 292, 196, 786, 72, "DASHBOARD SUMMARY HERO", "Forecast refresh, report shortcut, and analytics status")
    for i, label in enumerate(["Predicted Revenue", "Forecasted Units", "Top Selling Items", "Upsell Opportunities"]):
        labeled_box(svg, 292 + i * 196, 292, 180, 70, label, "KPI value")
    labeled_box(svg, 292, 386, 470, 140, "ITEM SALES TREND CHART", "Monthly product sales trend")
    labeled_box(svg, 786, 386, 292, 140, "ITEM HIGHLIGHTS", "Top product and forecast notes")
    labeled_box(svg, 292, 548, 370, 58, "RECENT TRANSACTIONS", "Latest sales/service activity")
    labeled_box(svg, 688, 548, 390, 58, "LOW STOCK ALERTS", "Predicted and current low stock")


def draw_inventory(svg: Svg):
    app_shell(svg, "Inventory Management")
    labeled_box(svg, 292, 196, 786, 60, "INVENTORY TOOLBAR", "Search products | Scan barcode | Add stock | Import price list | Print labels")
    labeled_box(svg, 292, 282, 520, 260, "PRODUCT RECORDS TABLE / GRID", "SKU | Part name | Category | Stock | Price | Status | Actions")
    labeled_box(svg, 838, 282, 240, 118, "FILTERS", "Category\nStock level\nCompatibility")
    labeled_box(svg, 838, 424, 240, 118, "PRODUCT DETAIL / LABEL PREVIEW", "Barcode label\nMitsubishi parts label\nUpdate stock")
    svg.rect(292, 566, 786, 40, "ADD STOCK / PRODUCT FORM MODAL AREA", cls="soft", size="small")


def draw_pos(svg: Svg):
    app_shell(svg, "Point of Sale")
    labeled_box(svg, 292, 196, 420, 60, "PRODUCT SEARCH AND BARCODE INPUT", "Scan or type SKU / part name")
    labeled_box(svg, 738, 196, 340, 60, "SALE CONTROLS", "New sale | Hold | Clear cart")
    labeled_box(svg, 292, 284, 420, 220, "PRODUCT RESULTS", "Item cards or table rows with add buttons")
    labeled_box(svg, 738, 284, 340, 220, "CART", "Selected items\nQuantity controls\nDiscounts")
    labeled_box(svg, 292, 532, 420, 74, "PAYMENT ENTRY", "Cash / GCash / bank transfer")
    labeled_box(svg, 738, 532, 340, 74, "RECEIPT PREVIEW", "Printable transaction receipt")


def draw_quotation(svg: Svg):
    app_shell(svg, "Cost Estimation & Quotation")
    labeled_box(svg, 292, 196, 370, 94, "CUSTOMER AND VEHICLE FORM", "Customer details\nVehicle model/year\nService concern")
    labeled_box(svg, 684, 196, 394, 94, "QUOTE ACTIONS", "Save quote | Print | Reset | Lookup")
    labeled_box(svg, 292, 318, 370, 188, "PARTS AND SERVICES SELECTOR", "Part search\nService selector\nRecommended bundles")
    labeled_box(svg, 684, 318, 394, 188, "QUOTATION LINE ITEMS", "Product / service\nQuantity\nPrice\nSubtotal")
    svg.rect(292, 532, 786, 72, "TOTAL COMPUTATION AND PRINTABLE QUOTATION OUTPUT", cls="soft", size="small")


def draw_services_admin(svg: Svg):
    app_shell(svg, "Service Orders")
    labeled_box(svg, 292, 196, 786, 56, "SERVICE ORDER STATUS FILTERS", "Pending | In progress | Completed | Cancelled")
    labeled_box(svg, 292, 278, 500, 250, "SERVICE ORDER LIST", "Customer | Vehicle | Concern | Assigned staff | Status | Date")
    labeled_box(svg, 818, 278, 260, 250, "ORDER DETAIL PANEL", "Diagnosis\nParts needed\nLabor/service notes\nUpdate progress")
    svg.rect(292, 556, 786, 48, "CREATE / ASSIGN / COMPLETE SERVICE ORDER MODALS", cls="soft", size="small")


def draw_stockroom(svg: Svg):
    app_shell(svg, "3D Stockroom")
    labeled_box(svg, 292, 196, 540, 330, "3D STOCKROOM VIEW", "Floor 1 sales area\nFloor 2 stockroom\nShelves, aisles, product markers")
    labeled_box(svg, 858, 196, 220, 86, "FLOOR CONTROLS", "Floor selector\nCamera controls")
    labeled_box(svg, 858, 306, 220, 96, "PRODUCT SEARCH", "Find part location")
    labeled_box(svg, 858, 426, 220, 100, "LOCATION DETAILS", "Shelf / level / slot\nRoute steps")
    svg.rect(292, 552, 786, 52, "ROUTE GUIDANCE AND STOCKROOM LEGEND", cls="soft", size="small")


def draw_parts_mapping(svg: Svg):
    app_shell(svg, "Parts Mapping Design")
    labeled_box(svg, 292, 196, 520, 330, "3D LAYOUT EDITOR CANVAS", "Edit floors, aisles, shelves, slots, and markers")
    labeled_box(svg, 838, 196, 240, 86, "EDITOR TOOLBAR", "Select | Move | Add shelf | Delete")
    labeled_box(svg, 838, 306, 240, 110, "STRUCTURE TREE", "Floors\nZones\nAisles\nShelves\nSlots")
    labeled_box(svg, 838, 440, 240, 86, "PROPERTIES PANEL", "Name\nPosition\nDimensions\nSave")
    svg.rect(292, 552, 786, 52, "SAVE LAYOUT / PUBLISH TO STOCKROOM LOCATOR", cls="soft", size="small")


def draw_reports(svg: Svg):
    app_shell(svg, "Reports & Analytics")
    labeled_box(svg, 292, 196, 786, 56, "REPORT TABS AND FILTERS", "Sales report | Inventory report | Date range | Export")
    for i, label in enumerate(["Total Sales", "Units Sold", "Low Stock", "Service Revenue"]):
        labeled_box(svg, 292 + i * 196, 278, 180, 62, label, "Summary card")
    labeled_box(svg, 292, 366, 470, 140, "ANALYTICS CHART", "Sales trend / inventory movement")
    labeled_box(svg, 786, 366, 292, 140, "REPORT SUMMARY PANEL", "Top products\nForecast notes\nStock risks")
    svg.rect(292, 532, 786, 70, "REPORT TABLE AND HISTORICAL SALE EDITOR MODAL", cls="soft", size="small")


def draw_users(svg: Svg):
    app_shell(svg, "User Management")
    for i, label in enumerate(["Total Users", "Active Staff", "Administrators"]):
        labeled_box(svg, 292 + i * 262, 196, 240, 64, label, "Count")
    labeled_box(svg, 292, 286, 786, 52, "SEARCH USERS AND ADD USER BUTTON", "Search users by name, email, or role")
    labeled_box(svg, 292, 364, 500, 162, "USER TABLE", "User | Email | Role | Status | Actions")
    labeled_box(svg, 818, 364, 260, 162, "ADD / EDIT USER MODAL", "Name\nEmail\nRole\nPassword")
    svg.rect(292, 554, 786, 44, "ROLE VALIDATION AND ACCOUNT MANAGEMENT", cls="soft", size="small")


def draw_cms(svg: Svg):
    app_shell(svg, "Content CMS")
    labeled_box(svg, 292, 196, 210, 330, "PAGE LIST", "Home\nAbout\nCatalog\nService Orders")
    labeled_box(svg, 528, 196, 314, 330, "SECTION EDITOR", "Hero content\nCards\nFAQ\nContact blocks")
    labeled_box(svg, 868, 196, 210, 330, "LIVE PREVIEW", "Public page preview")
    svg.rect(292, 552, 786, 52, "SAVE DRAFT / PUBLISH CONTENT / REORDER SECTIONS", cls="soft", size="small")


DRAWERS = {
    "public_home": draw_public_home,
    "public_catalog": draw_public_catalog,
    "estimate_choice": draw_estimate_choice,
    "estimate_builder": draw_estimate_builder,
    "quote_summary": draw_quote_summary,
    "public_service_orders": draw_public_service_orders,
    "public_about": draw_public_about,
    "login": draw_login,
    "workspace_shell": draw_workspace_shell,
    "dashboard": draw_dashboard,
    "inventory": draw_inventory,
    "pos": draw_pos,
    "quotation": draw_quotation,
    "services_admin": draw_services_admin,
    "stockroom": draw_stockroom,
    "parts_mapping": draw_parts_mapping,
    "reports": draw_reports,
    "users": draw_users,
    "cms": draw_cms,
}


def generate_svgs():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    manifest = []
    for index, fig in enumerate(FIGURES, start=1):
        number = f"{index:02d}"
        filename = f"figure-{number}-{fig['slug']}"
        title = f"Figure {index}. {fig['title']}"
        svg = Svg(title)
        DRAWERS[fig["layout"]](svg)
        svg_path = FIG_DIR / f"{filename}.svg"
        svg_path.write_text(svg.finish(), encoding="utf-8")
        manifest.append(
            {
                "number": index,
                "title": fig["title"],
                "description": fig["description"].format(n=index),
                "svg": str(svg_path),
                "png": str(FIG_DIR / f"{filename}.png"),
            }
        )
    (OUT_DIR / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    lines = ["# LimenServe System Wireframe Figure Descriptions", ""]
    for item in manifest:
        lines.append(f"## Figure {item['number']}. {item['title']}")
        lines.append(item["description"])
        lines.append("")
    (OUT_DIR / "figure-descriptions.md").write_text("\n".join(lines), encoding="utf-8")
    return manifest


def build_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    manifest_path = OUT_DIR / "manifest.json"
    if not manifest_path.exists():
        raise SystemExit("Run the svg command first.")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    missing = [item["png"] for item in manifest if not Path(item["png"]).exists()]
    if missing:
        raise SystemExit("Missing rendered PNG files:\n" + "\n".join(missing))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LIMENSERVE SYSTEM WIREFRAMES")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "The following black-and-white wireframes present the main public and internal interfaces of LimenServe. "
        "Each figure uses labeled layout blocks to show the intended structure, navigation, and major controls of the system."
    )

    for item in manifest:
        doc.add_paragraph()
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(item["png"], width=Inches(6.85))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cap.add_run(f"Figure {item['number']}. ")
        r1.bold = True
        r2 = cap.add_run(item["title"])
        r1.font.name = r2.font.name = "Times New Roman"
        r1.font.size = r2.font.size = Pt(11)

        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc.paragraph_format.first_line_indent = Inches(0.25)
        desc.add_run(item["description"])

        if item["number"] != manifest[-1]["number"]:
            doc.add_page_break()

    doc_path = OUT_DIR / "LimenServe_System_Wireframes_Black_White.docx"
    doc.save(doc_path)
    return doc_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["svg", "docx", "all"])
    args = parser.parse_args()

    if args.command in {"svg", "all"}:
        generate_svgs()
    if args.command in {"docx", "all"}:
        path = build_docx()
        print(path)


if __name__ == "__main__":
    main()
