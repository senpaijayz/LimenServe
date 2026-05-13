from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Methodology_With_Diagrams.docx"
RRL_MD = ROOT / "Docs" / "Refined_RRL_LimenServe.md"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Methodology_With_Diagrams_RRL_Refined.docx"


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def parse_refined_rrl(md_text: str):
    blocks = []
    current = []

    def flush():
        nonlocal current
        if current:
            blocks.append(("p", " ".join(line.strip() for line in current).strip()))
            current = []

    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("# "):
            flush()
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            flush()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            flush()
            blocks.append(("h3", line[4:].strip()))
        else:
            current.append(line)
    flush()
    return blocks


def format_paragraph(paragraph, kind: str):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        if kind == "h1":
            run.font.size = Pt(12)
            run.bold = True
        elif kind == "h2":
            run.font.size = Pt(12)
            run.bold = True
        elif kind == "h3":
            run.font.size = Pt(11)
            run.bold = True
        else:
            run.font.size = Pt(11)
    if kind == "h1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    elif kind == "h2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    elif kind == "h3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(36)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)


def main():
    doc = Document(SOURCE_DOCX)
    start = None
    end = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().upper()
        if text == "REVIEW OF RELATED LITERATURE" and start is None:
            start = index
        elif start is not None and text == "METHODOLOGY":
            end = index
            break
    if start is None or end is None:
        raise SystemExit("Could not locate RRL or METHODOLOGY boundary.")

    anchor = doc.paragraphs[start - 1]
    for paragraph in list(doc.paragraphs[start:end]):
        delete_paragraph(paragraph)

    cursor = insert_paragraph_after(anchor)
    cursor.add_run().add_break(WD_BREAK.PAGE)
    for kind, text in parse_refined_rrl(RRL_MD.read_text(encoding="utf-8")):
        cursor = insert_paragraph_after(cursor, text)
        format_paragraph(cursor, kind)

    doc.save(OUT_DOCX)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().upper() == "METHODOLOGY":
            paragraph.paragraph_format.page_break_before = True
            break
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
