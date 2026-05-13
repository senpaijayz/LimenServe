from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Feasibility_Gantt.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Appendix_Gantt.docx"
GANTT_IMAGE = ROOT / "Docs" / "limenserve_appendix_gantt.png"

RESEARCHERS = (
    "Jay Iverson A. Dela Cruz, Alden M. Suacillo, Don Rich M. Ulanday, "
    "Christian Dave D. Vargas, Estefhano C. Villafuerte"
)
PROGRAM = "Bachelor of Science in Information Technology"
TITLE = (
    "LimenServe: A Web-Based Management Information System with Automated Cost "
    "Estimation and Quotation for Services and Orders of Limen Auto Parts Center"
)

MONTHS = [
    ("2025", "Oct"),
    ("2025", "Nov"),
    ("2025", "Dec"),
    ("2026", "Jan"),
    ("2026", "Feb"),
    ("2026", "Mar"),
    ("2026", "Apr"),
    ("2026", "May"),
    ("2026", "Jun"),
    ("2026", "Jul"),
    ("2026", "Aug"),
    ("2026", "Sep"),
    ("2026", "Oct"),
    ("2026", "Nov"),
    ("2026", "Dec"),
    ("2027", "Jan"),
    ("2027", "Feb"),
]

SCHEDULE = [
    ("Proposal drafting", 0, 2),
    ("Proposal defense", 3, 3),
    ("System design", 4, 5),
    ("System development", 5, 14),
    ("System testing", 12, 14),
    ("System evaluation", 13, 14),
    ("System implementation", 13, 14),
    ("Manuscript preparation", 0, 15),
    ("Final oral defense", 15, 15),
    ("Manuscript review", 14, 16),
]


def insert_paragraph_before(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def remove_inline_gantt(doc: Document):
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "Gantt Chart. LimenServe development schedule":
            previous = paragraph._p.getprevious()
            if previous is not None and previous.tag == qn("w:tbl"):
                previous.getparent().remove(previous)
            paragraph._p.getparent().remove(paragraph._p)
            break


def update_schedule_reference(doc: Document):
    old = "The Gantt chart below presents the major activities and target schedule for the development of LimenServe."
    new = "The Gantt chart in Appendix 3 presents the major activities and target schedule for the development of LimenServe."
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(old):
            paragraph.text = paragraph.text.replace(old, new)
            format_paragraph(paragraph, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            return


def font(size: int, bold: bool = False, italic: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbi.ttf" if bold and italic else None,
        "C:/Windows/Fonts/arialbd.ttf" if bold else None,
        "C:/Windows/Fonts/ariali.ttf" if italic else None,
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def multiline_center(draw, xy, text, fnt, fill=(30, 30, 30), spacing=5):
    lines = text.split("\n")
    widths = [draw.textbbox((0, 0), line, font=fnt)[2] for line in lines]
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] - draw.textbbox((0, 0), line, font=fnt)[1] for line in lines]
    x, y = xy
    for line, width, height in zip(lines, widths, heights):
        draw.text((x - width / 2, y), line, font=fnt, fill=fill)
        y += height + spacing


def wrap_text(draw, text: str, fnt, max_width: int):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        probe = f"{current} {word}".strip()
        if draw.textbbox((0, 0), probe, font=fnt)[2] <= max_width or not current:
            current = probe
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, box, text, fnt, fill=(30, 30, 30), align="left", spacing=4):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1)
    line_heights = [draw.textbbox((0, 0), line, font=fnt)[3] - draw.textbbox((0, 0), line, font=fnt)[1] for line in lines]
    total = sum(line_heights) + spacing * (len(lines) - 1)
    y = y1 + max(0, (y2 - y1 - total) / 2)
    for line, height in zip(lines, line_heights):
        width = draw.textbbox((0, 0), line, font=fnt)[2]
        x = x1 if align == "left" else x1 + (x2 - x1 - width) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += height + spacing


def generate_gantt_image():
    img = Image.new("RGB", (1700, 2200), "white")
    d = ImageDraw.Draw(img)
    black = (35, 35, 35)
    grid = (25, 25, 25)
    green = (112, 145, 115)
    header_fill = (232, 232, 232)

    d.text((135, 70), "Appendix 3. Gantt Chart", font=font(32, bold=True), fill=black)
    multiline_center(d, (850, 170), "Republic of the Philippines", font(30), black)
    multiline_center(d, (850, 225), "CAVITE STATE UNIVERSITY", font(46, bold=True), black)
    multiline_center(d, (850, 295), "Bacoor City Campus", font(33, bold=True), black)
    multiline_center(d, (850, 345), "SHIV, Molino VI, City of Bacoor", font(29), black)
    multiline_center(d, (850, 420), "DEPARTMENT OF COMPUTER STUDIES", font(34, bold=True), black)
    multiline_center(d, (850, 495), "GANTT CHART", font(36, bold=True), black)

    info_x, info_y = 230, 585
    label_w, colon_w, value_w = 360, 50, 860
    row_h = 62
    rows = [
        ("Name of Researcher(s)", RESEARCHERS),
        ("Program", PROGRAM),
        ("Title of Study", TITLE),
    ]
    for idx, (label, value) in enumerate(rows):
        y1 = info_y + idx * row_h
        y2 = y1 + row_h
        d.text((info_x, y1 + 18), label, font=font(25, italic=True), fill=black)
        d.text((info_x + label_w + 15, y1 + 18), ":", font=font(25), fill=black)
        d.line((info_x + label_w + colon_w, y2 - 9, info_x + label_w + colon_w + value_w, y2 - 9), fill=black, width=2)
        draw_wrapped(d, (info_x + label_w + colon_w + 10, y1, info_x + label_w + colon_w + value_w, y2 - 8), value, font(24), black)

    table_x, table_y = 115, 835
    act_w, month_w = 235, 75
    h_year, h_month, row_h = 50, 62, 78
    table_w = act_w + month_w * len(MONTHS)
    table_h = h_year + h_month + row_h * len(SCHEDULE)
    d.rectangle((table_x, table_y, table_x + table_w, table_y + table_h), outline=grid, width=2)
    d.rectangle((table_x, table_y, table_x + table_w, table_y + h_year + h_month), fill=header_fill, outline=grid, width=2)

    d.rectangle((table_x, table_y, table_x + act_w, table_y + h_year + h_month), outline=grid, width=2)
    multiline_center(d, (table_x + act_w / 2, table_y + 38), "ACTIVITY", font(25, bold=True), black)

    year_ranges = [("2025", 0, 2), ("2026", 3, 14), ("2027", 15, 16)]
    for year, start, end in year_ranges:
        x1 = table_x + act_w + start * month_w
        x2 = table_x + act_w + (end + 1) * month_w
        d.rectangle((x1, table_y, x2, table_y + h_year), outline=grid, width=2)
        multiline_center(d, ((x1 + x2) / 2, table_y + 11), year, font(24, bold=True), black)

    for idx, (_, month) in enumerate(MONTHS):
        x1 = table_x + act_w + idx * month_w
        x2 = x1 + month_w
        d.rectangle((x1, table_y + h_year, x2, table_y + h_year + h_month), outline=grid, width=2)
        multiline_center(d, ((x1 + x2) / 2, table_y + h_year + 12), month, font(18, bold=True), black)

    for row_idx, (activity, start, end) in enumerate(SCHEDULE):
        y1 = table_y + h_year + h_month + row_idx * row_h
        y2 = y1 + row_h
        d.rectangle((table_x, y1, table_x + act_w, y2), outline=grid, width=2)
        draw_wrapped(d, (table_x + 12, y1 + 4, table_x + act_w - 12, y2 - 4), activity, font(18, italic=True), black, align="center", spacing=2)
        for month_idx in range(len(MONTHS)):
            x1 = table_x + act_w + month_idx * month_w
            x2 = x1 + month_w
            fill = green if start <= month_idx <= end else "white"
            d.rectangle((x1, y1, x2, y2), fill=fill, outline=grid, width=2)

    caption = "Figure A1. Gantt Chart of LimenServe development schedule"
    w = d.textbbox((0, 0), caption, font=font(28, italic=True))[2]
    d.text(((1700 - w) / 2, table_y + table_h + 24), caption, font=font(28, italic=True), fill=black)

    img.save(GANTT_IMAGE, quality=95)


def add_appendix_gantt(doc: Document):
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Appendix A. Interview Report":
            anchor = paragraph
            break
    if anchor is None:
        raise SystemExit("Could not find Appendix A insertion point.")

    generate_gantt_image()

    top = insert_paragraph_before(anchor, "")
    top.paragraph_format.page_break_before = True
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top.add_run()
    run.add_picture(str(GANTT_IMAGE), width=Inches(6.55))


def main():
    doc = Document(SOURCE_DOCX)
    remove_inline_gantt(doc)
    update_schedule_reference(doc)
    if any(p.text.strip() == "Appendix 3. Gantt Chart" for p in doc.paragraphs):
        raise SystemExit("Appendix 3 Gantt chart already exists.")
    add_appendix_gantt(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
