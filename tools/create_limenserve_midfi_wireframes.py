from __future__ import annotations

import argparse
import html
import json
from pathlib import Path
from textwrap import wrap


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docs" / "wireframes" / "limenserve-system-wireframes-midfi"
FIG_DIR = OUT_DIR / "figures"
WIDTH = 1200
HEIGHT = 720


COLORS = {
    "bg": "#f8f9fa",
    "surface": "#ffffff",
    "surface2": "#f1f5f9",
    "border": "#e5e7eb",
    "border2": "#cbd5e1",
    "text": "#111827",
    "muted": "#64748b",
    "soft": "#e2e8f0",
    "dark": "#0f172a",
    "blue": "#1E3A8A",
    "blue2": "#1D4ED8",
    "blueSoft": "#eff6ff",
    "red": "#DC2626",
    "redSoft": "#fee2e2",
    "green": "#10b981",
    "greenSoft": "#d1fae5",
    "amber": "#f59e0b",
    "amberSoft": "#fef3c7",
    "purple": "#7c3aed",
    "purpleSoft": "#ede9fe",
}


FIGURES = [
    {
        "slug": "public-home",
        "title": "Public Home Page Mid-Fidelity Wireframe",
        "layout": "public_home",
        "description": (
            "Figure {n} shows the public landing page of LimenServe. It presents the shop identity, "
            "main navigation, search entry point, featured vehicle categories, product highlights, "
            "trust indicators, and call-to-action areas that guide customers to browse parts or request an estimate."
        ),
    },
    {
        "slug": "public-catalog",
        "title": "Public Parts Catalog Mid-Fidelity Wireframe",
        "layout": "public_catalog",
        "description": (
            "Figure {n} shows the public parts catalog layout. It includes vehicle-first filters, search controls, "
            "product cards, matched service package suggestions, and a product detail area for customer browsing."
        ),
    },
    {
        "slug": "public-estimate-choice",
        "title": "Public Estimate Entry Mid-Fidelity Wireframe",
        "layout": "estimate_choice",
        "description": (
            "Figure {n} shows the first phase of the public estimate workflow. Customers choose whether to build "
            "a new estimate or retrieve an existing quotation, then provide customer and vehicle details before "
            "continuing to parts and services."
        ),
    },
    {
        "slug": "public-estimate-builder",
        "title": "Public Estimate Builder Mid-Fidelity Wireframe",
        "layout": "estimate_builder",
        "description": (
            "Figure {n} shows the public estimate builder. The layout combines part search, service selection, "
            "smart bundle recommendations, and a quotation cart so customers can review selected parts and services "
            "before finalizing the estimate."
        ),
    },
    {
        "slug": "public-quote-summary",
        "title": "Public Quotation Summary Mid-Fidelity Wireframe",
        "layout": "quote_summary",
        "description": (
            "Figure {n} shows the public quotation summary and print preview structure. It presents customer and "
            "vehicle information, quotation line items, service lines, subtotals, estimated total, and printing controls."
        ),
    },
    {
        "slug": "public-service-orders",
        "title": "Public Service Orders Mid-Fidelity Wireframe",
        "layout": "public_service_orders",
        "description": (
            "Figure {n} shows the public service-order information page. It explains the service workflow, from "
            "initial assessment to quotation, service handling, status tracking, and completion."
        ),
    },
    {
        "slug": "public-about",
        "title": "Public About Page Mid-Fidelity Wireframe",
        "layout": "public_about",
        "description": (
            "Figure {n} shows the public about page of Limen Auto Parts Center. It contains the business story, "
            "operational overview, shop statistics, mechanics section, service values, and location information."
        ),
    },
    {
        "slug": "login",
        "title": "Staff Login Page Mid-Fidelity Wireframe",
        "layout": "login",
        "description": (
            "Figure {n} shows the staff login page. It provides controlled access to the internal workspace through "
            "email and password fields, remember-me option, sign-in action, and administrative contact notice."
        ),
    },
    {
        "slug": "workspace-shell",
        "title": "Internal Workspace Shell Mid-Fidelity Wireframe",
        "layout": "workspace_shell",
        "description": (
            "Figure {n} shows the common internal workspace shell used after authentication. It includes the sidebar "
            "navigation, top header, global search area, notification area, user identity section, and protected content area."
        ),
    },
    {
        "slug": "dashboard",
        "title": "Admin Dashboard Mid-Fidelity Wireframe",
        "layout": "dashboard",
        "description": (
            "Figure {n} shows the dashboard module. It summarizes operational indicators through KPI cards, forecast "
            "widgets, sales trend charts, item highlights, recent transactions, and low-stock alerts."
        ),
    },
    {
        "slug": "inventory",
        "title": "Inventory Management Mid-Fidelity Wireframe",
        "layout": "inventory",
        "description": (
            "Figure {n} shows the inventory management module. It provides search and filter controls, product listings, "
            "stock status indicators, barcode and label support, price-list import, and add-stock actions for inventory updates."
        ),
    },
    {
        "slug": "pos",
        "title": "Point of Sale Mid-Fidelity Wireframe",
        "layout": "pos",
        "description": (
            "Figure {n} shows the point-of-sale module. It supports barcode or product search, cart building, quantity "
            "updates, payment entry, transaction total computation, and receipt preview for completed sales."
        ),
    },
    {
        "slug": "quotation",
        "title": "Internal Quotation Builder Mid-Fidelity Wireframe",
        "layout": "quotation",
        "description": (
            "Figure {n} shows the internal quotation builder used by staff. It combines customer details, vehicle information, "
            "part selection, service selection, quotation line items, total computation, and print or save actions."
        ),
    },
    {
        "slug": "service-orders-admin",
        "title": "Service Orders Management Mid-Fidelity Wireframe",
        "layout": "services_admin",
        "description": (
            "Figure {n} shows the internal service-order management module. It organizes service requests by status, "
            "supports search and filtering, displays order details, and allows staff to assign personnel or update progress."
        ),
    },
    {
        "slug": "stockroom",
        "title": "3D Stockroom Locator Mid-Fidelity Wireframe",
        "layout": "stockroom",
        "description": (
            "Figure {n} shows the stockroom visualization module. It includes a 3D stockroom area, floor controls, "
            "product search, item location details, and route steps that help staff locate inventory inside the shop."
        ),
    },
    {
        "slug": "parts-mapping",
        "title": "Parts Mapping Admin Mid-Fidelity Wireframe",
        "layout": "parts_mapping",
        "description": (
            "Figure {n} shows the parts mapping administration interface. It provides tools for editing the stockroom layout, "
            "mapping shelves and slots, modifying object properties, and saving layout changes used by the locator module."
        ),
    },
    {
        "slug": "reports",
        "title": "Reports and Analytics Mid-Fidelity Wireframe",
        "layout": "reports",
        "description": (
            "Figure {n} shows the reporting and analytics module. It contains report tabs, date filters, summary cards, "
            "sales and inventory charts, tabular report results, and tools for encoding or editing historical sales data."
        ),
    },
    {
        "slug": "users",
        "title": "User Management Mid-Fidelity Wireframe",
        "layout": "users",
        "description": (
            "Figure {n} shows the user management module. It displays user summary cards, search controls, role-based user "
            "records, and an add or edit user modal for maintaining authorized staff accounts."
        ),
    },
    {
        "slug": "cms",
        "title": "Content CMS Mid-Fidelity Wireframe",
        "layout": "cms",
        "description": (
            "Figure {n} shows the content management module. It allows administrators to manage public pages, edit page sections, "
            "preview content, and publish updates that appear on the public LimenServe website."
        ),
    },
]


def esc(value: object) -> str:
    return html.escape(str(value), quote=True)


class Svg:
    def __init__(self):
        self.parts: list[str] = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{WIDTH}" height="{HEIGHT}" viewBox="0 0 {WIDTH} {HEIGHT}">',
            "<defs>",
            '<linearGradient id="heroGrad" x1="0" y1="0" x2="1" y2="1"><stop offset="0" stop-color="#eff6ff"/><stop offset="0.55" stop-color="#ffffff"/><stop offset="1" stop-color="#fee2e2"/></linearGradient>',
            '<linearGradient id="darkGrad" x1="0" y1="0" x2="1" y2="1"><stop offset="0" stop-color="#111827"/><stop offset="1" stop-color="#1E3A8A"/></linearGradient>',
            '<filter id="shadow" x="-10%" y="-10%" width="120%" height="140%"><feDropShadow dx="0" dy="12" stdDeviation="12" flood-color="#0f172a" flood-opacity="0.10"/></filter>',
            "</defs>",
            "<style>",
            "text{font-family:Inter,Arial,Helvetica,sans-serif;fill:#111827;letter-spacing:0}",
            ".h1{font-size:34px;font-weight:800}",
            ".h2{font-size:22px;font-weight:800}",
            ".h3{font-size:15px;font-weight:800}",
            ".body{font-size:12px;font-weight:500}",
            ".small{font-size:10px;font-weight:600}",
            ".tiny{font-size:8px;font-weight:700}",
            ".muted{fill:#64748b}",
            ".white{fill:#ffffff}",
            ".blue{fill:#1E3A8A}",
            ".red{fill:#DC2626}",
            "</style>",
            f'<rect x="0" y="0" width="{WIDTH}" height="{HEIGHT}" fill="{COLORS["bg"]}"/>',
        ]

    def rect(self, x, y, w, h, fill="surface", stroke="border", rx=16, opacity=None, shadow=False):
        fill_value = COLORS.get(fill, fill)
        stroke_value = COLORS.get(stroke, stroke)
        filter_value = ' filter="url(#shadow)"' if shadow else ""
        opacity_value = f' opacity="{opacity}"' if opacity is not None else ""
        self.parts.append(
            f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill_value}" stroke="{stroke_value}" stroke-width="1"{filter_value}{opacity_value}/>'
        )

    def text(self, x, y, value, cls="body", anchor="start", color=None):
        color_style = f' style="fill:{COLORS.get(color, color)}"' if color else ""
        self.parts.append(f'<text x="{x}" y="{y}" text-anchor="{anchor}" class="{cls}"{color_style}>{esc(value)}</text>')

    def multiline(self, value, x, y, width, cls="body", color=None, line_height=15, anchor="start"):
        lines: list[str] = []
        for paragraph in str(value).split("\n"):
            lines.extend(wrap(paragraph, width=max(8, int(width / 7.0))) or [""])
        for i, line in enumerate(lines):
            self.text(x, y + i * line_height, line, cls=cls, anchor=anchor, color=color)

    def circle(self, cx, cy, r, fill="surface2", stroke="border"):
        self.parts.append(
            f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="{COLORS.get(fill, fill)}" stroke="{COLORS.get(stroke, stroke)}" stroke-width="1"/>'
        )

    def line(self, x1, y1, x2, y2, color="border2", width=2):
        self.parts.append(
            f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{COLORS.get(color, color)}" stroke-width="{width}" stroke-linecap="round"/>'
        )

    def path(self, d, color="blue", width=3, fill="none"):
        self.parts.append(
            f'<path d="{d}" fill="{fill}" stroke="{COLORS.get(color, color)}" stroke-width="{width}" stroke-linecap="round" stroke-linejoin="round"/>'
        )

    def pill(self, x, y, w, h, label, fill="surface2", color="text"):
        self.rect(x, y, w, h, fill=fill, stroke="border", rx=h / 2)
        self.text(x + w / 2, y + h / 2 + 4, label, cls="small", anchor="middle", color=color)

    def button(self, x, y, w, h, label, primary=False):
        self.rect(x, y, w, h, fill="blue" if primary else "surface", stroke="blue" if primary else "border2", rx=12)
        self.text(x + w / 2, y + h / 2 + 4, label, cls="small", anchor="middle", color="surface" if primary else "text")

    def input(self, x, y, w, h, label):
        self.rect(x, y, w, h, fill="surface", stroke="border2", rx=12)
        self.text(x + 14, y + h / 2 + 4, label, cls="small", color="muted")

    def table(self, x, y, w, h, headers, rows=4):
        self.rect(x, y, w, h, fill="surface", stroke="border", rx=16, shadow=True)
        header_h = 34
        self.parts.append(f'<path d="M{x} {y + header_h} H{x + w}" stroke="{COLORS["border"]}" stroke-width="1"/>')
        col_w = w / len(headers)
        for i, header in enumerate(headers):
            self.text(x + 16 + i * col_w, y + 22, header, cls="tiny", color="muted")
        row_h = (h - header_h) / rows
        for r in range(rows):
            yy = y + header_h + r * row_h
            if r > 0:
                self.line(x + 14, yy, x + w - 14, yy, color="border", width=1)
            for i in range(len(headers)):
                self.rect(x + 16 + i * col_w, yy + 12, min(70, col_w - 26), 8, fill="soft", stroke="soft", rx=4)

    def finish(self):
        self.parts.append("</svg>")
        return "\n".join(self.parts)


def logo(svg: Svg, x, y, dark=False):
    svg.rect(x, y, 42, 42, fill="surface", stroke="border", rx=10)
    svg.text(x + 21, y + 27, "L", cls="h2", anchor="middle", color="red")
    svg.text(x + 52, y + 17, "LIMEN", cls="h3", color="surface" if dark else "text")
    svg.text(x + 52, y + 33, "Genuine Auto Parts", cls="small", color="surface" if dark else "muted")


def public_shell(svg: Svg, active="Home"):
    svg.rect(0, 0, WIDTH, 34, fill="dark", stroke="dark", rx=0)
    svg.text(64, 22, "Genuine and aftermarket parts from a real Pasay City store", cls="small", color="surface")
    svg.text(850, 22, "(0915) 522 5629 | Mon-Sat 8:00 AM-5:00 PM", cls="small", color="surface")
    svg.rect(0, 34, WIDTH, 78, fill="surface", stroke="border", rx=0, shadow=True)
    logo(svg, 64, 52)
    nav = ["Home", "About", "Genuine Parts", "Get Estimate", "Service Orders", "Staff Portal"]
    x = 420
    for item in nav:
        w = 70 if item in {"Home", "About"} else 108 if item == "Genuine Parts" else 104
        is_active = item == active
        svg.pill(x, 60, w, 30, item, fill="blueSoft" if is_active else "surface", color="blue" if is_active else "text")
        x += w + 12


def internal_shell(svg: Svg, active="Dashboard"):
    svg.rect(0, 0, WIDTH, HEIGHT, fill="bg", stroke="bg", rx=0)
    svg.rect(0, 0, 250, HEIGHT, fill="surface", stroke="border", rx=0, shadow=True)
    logo(svg, 24, 22)
    svg.text(24, 92, "MAIN MENU", cls="tiny", color="muted")
    nav = ["Dashboard", "Point of Sale", "Inventory", "Quotation", "Service Orders", "3D Stockroom", "Reports", "User Management", "Content CMS"]
    y = 112
    for item in nav:
        is_active = item == active
        svg.rect(20, y, 210, 36, fill="blueSoft" if is_active else "surface", stroke="blueSoft" if is_active else "surface", rx=10)
        if is_active:
            svg.rect(20, y + 6, 4, 24, fill="blue", stroke="blue", rx=2)
        svg.circle(42, y + 18, 8, fill="blue" if is_active else "surface2", stroke="blue" if is_active else "border")
        svg.text(60, y + 23, item, cls="small", color="blue" if is_active else "text")
        y += 42
    svg.rect(20, 642, 210, 54, fill="surface2", stroke="border", rx=14)
    svg.circle(44, 669, 15, fill="redSoft", stroke="redSoft")
    svg.text(68, 664, "Admin User", cls="small")
    svg.text(68, 680, "Administrator", cls="tiny", color="muted")
    svg.rect(250, 0, 950, 72, fill="surface", stroke="border", rx=0)
    svg.text(282, 43, active, cls="h2")
    svg.input(730, 20, 230, 34, "Search modules, products...")
    svg.button(974, 20, 76, 34, "Alerts")
    svg.button(1062, 20, 92, 34, "Profile")


def page_title(svg: Svg, x, y, eyebrow, title, body):
    svg.text(x, y, eyebrow.upper(), cls="tiny", color="blue")
    svg.text(x, y + 36, title, cls="h1")
    svg.multiline(body, x, y + 62, 470, cls="body", color="muted", line_height=16)


def metric(svg: Svg, x, y, w, label, value, fill="surface"):
    svg.rect(x, y, w, 78, fill=fill, stroke="border", rx=16, shadow=True)
    svg.text(x + 18, y + 25, label, cls="small", color="muted")
    svg.text(x + 18, y + 55, value, cls="h2")


def product_card(svg: Svg, x, y, w=198, h=126, label="Product", price="PHP 1,250"):
    svg.rect(x, y, w, h, fill="surface", stroke="border", rx=16, shadow=True)
    svg.rect(x + 12, y + 12, 56, 42, fill="surface2", stroke="border", rx=10)
    svg.text(x + 80, y + 28, label, cls="h3")
    svg.text(x + 80, y + 45, "Mitsubishi compatible", cls="tiny", color="muted")
    svg.text(x + 16, y + 82, price, cls="h3", color="blue")
    svg.pill(x + 16, y + 94, 72, 22, "In stock", fill="greenSoft", color="green")
    svg.button(x + w - 82, y + 88, 66, 26, "Details")


def draw_public_home(svg: Svg):
    public_shell(svg, "Home")
    svg.rect(64, 140, 1072, 240, fill="url(#heroGrad)", stroke="border", rx=28, shadow=True)
    page_title(svg, 96, 182, "Genuine Parts", "Limen Auto Parts Center", "Search by part name, part number, or vehicle model and move straight into a cleaner quotation flow.")
    svg.input(96, 290, 360, 44, "Search by part, vehicle, or part number")
    svg.button(468, 290, 132, 44, "Search Parts", primary=True)
    for i, label in enumerate(["Montero Sport", "Triton", "Xforce", "Xpander"]):
        svg.pill(96 + i * 116, 346, 104, 28, label, fill="surface")
    svg.rect(676, 166, 390, 184, fill="darkGrad", stroke="dark", rx=24)
    svg.text(718, 236, "Storefront / Vehicle Showcase", cls="h2", color="surface")
    svg.text(718, 264, "Photo area in the actual system", cls="body", color="surface")
    for i, (label, value) in enumerate([("Catalog", "Search by part"), ("Estimate", "Fast quotation"), ("Orders", "Service workflow")]):
        metric(svg, 64 + i * 366, 408, 330, label, value, fill="surface")
    for i, label in enumerate(["Filters and fluids", "Electrical parts", "Service bundles", "Trust signals"]):
        product_card(svg, 64 + i * 268, 520, 238, 120, label, "Featured")


def draw_public_catalog(svg: Svg):
    public_shell(svg, "Genuine Parts")
    page_title(svg, 64, 148, "Parts Catalog", "Trusted auto parts", "Vehicle-first filtering helps customers browse cleaner catalog results.")
    svg.rect(650, 132, 486, 120, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(680, 166, "Choose your vehicle before you browse", cls="h3")
    svg.input(680, 184, 128, 36, "Model")
    svg.input(822, 184, 106, 36, "Year")
    svg.input(942, 184, 148, 36, "Variant / Engine")
    svg.input(64, 276, 680, 42, "Search by part name, part number, or vehicle")
    svg.button(760, 276, 126, 42, "Reset Filters")
    svg.button(900, 276, 120, 42, "Estimate", primary=True)
    for i, label in enumerate(["Oil Filter", "Brake Pad", "Air Filter", "Spark Plug", "Shock Absorber", "Cabin Filter"]):
        product_card(svg, 64 + (i % 3) * 236, 344 + (i // 3) * 146, 214, 126, label, f"PHP {850 + i * 220:,}")
    svg.rect(798, 344, 338, 270, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(826, 380, "Selected product", cls="h2")
    svg.rect(826, 402, 120, 82, fill="surface2", stroke="border", rx=16)
    svg.text(970, 426, "Compatibility", cls="h3")
    svg.multiline("Model match, stock status, SKU, and bundle suggestions appear here.", 970, 448, 130, cls="small", color="muted")
    svg.button(826, 520, 132, 38, "Add to Estimate", primary=True)
    svg.button(970, 520, 116, 38, "View Details")


def draw_estimate_choice(svg: Svg):
    public_shell(svg, "Get Estimate")
    page_title(svg, 64, 148, "Quotations", "Get Estimate", "Start a customer quotation or retrieve an existing quote number.")
    svg.rect(64, 286, 500, 142, fill="surface", stroke="border", rx=22, shadow=True)
    svg.circle(104, 330, 24, fill="blueSoft", stroke="blueSoft")
    svg.text(142, 324, "Build Estimate", cls="h2")
    svg.multiline("Pick parts, add services, and build a draft quotation.", 142, 350, 330, cls="body", color="muted")
    svg.button(142, 382, 136, 34, "Start Estimate", primary=True)
    svg.rect(596, 286, 500, 142, fill="surface", stroke="border", rx=22, shadow=True)
    svg.circle(636, 330, 24, fill="redSoft", stroke="redSoft")
    svg.text(674, 324, "Retrieve Quote", cls="h2")
    svg.multiline("Find a saved quotation and open a printable preview.", 674, 350, 330, cls="body", color="muted")
    svg.button(674, 382, 136, 34, "Lookup Quote")
    svg.rect(64, 460, 500, 148, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(94, 494, "Customer Details", cls="h3")
    svg.input(94, 512, 200, 36, "Full name")
    svg.input(310, 512, 200, 36, "Phone number")
    svg.input(94, 560, 416, 36, "Email address")
    svg.rect(596, 460, 500, 148, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(626, 494, "Vehicle Details", cls="h3")
    svg.input(626, 512, 144, 36, "Model")
    svg.input(786, 512, 110, 36, "Year")
    svg.input(912, 512, 140, 36, "Plate")
    svg.button(626, 560, 184, 36, "Continue", primary=True)


def draw_estimate_builder(svg: Svg):
    public_shell(svg, "Get Estimate")
    svg.text(64, 150, "Parts and Services", cls="h1")
    svg.text(64, 176, "Phase 2 cart with vehicle-matched products and service lines.", cls="body", color="muted")
    svg.input(64, 206, 440, 40, "Search parts for selected vehicle")
    svg.button(520, 206, 118, 40, "Sort")
    svg.rect(64, 272, 325, 318, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(90, 306, "Genuine Parts", cls="h2")
    for i, label in enumerate(["Oil Filter", "Brake Pad", "Cabin Filter"]):
        product_card(svg, 90, 326 + i * 78, 250, 64, label, f"PHP {900 + i * 600:,}")
    svg.rect(418, 272, 325, 318, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(444, 306, "Service Catalog", cls="h2")
    for i, label in enumerate(["Change Oil Service", "Brake Inspection", "General Checkup"]):
        svg.rect(444, 326 + i * 78, 250, 56, fill="surface2", stroke="border", rx=14)
        svg.text(462, 350 + i * 78, label, cls="h3")
        svg.text(462, 368 + i * 78, "Labor line", cls="tiny", color="muted")
    svg.rect(772, 142, 364, 448, fill="surface", stroke="border", rx=24, shadow=True)
    svg.text(804, 184, "Your quotation", cls="h2")
    svg.text(804, 206, "Selected parts and services", cls="body", color="muted")
    svg.table(804, 232, 300, 184, ["Item", "Qty", "Total"], rows=4)
    svg.rect(804, 440, 300, 74, fill="blueSoft", stroke="blueSoft", rx=18)
    svg.text(824, 468, "Estimated total", cls="small", color="blue")
    svg.text(824, 498, "PHP 8,450", cls="h2", color="blue")
    svg.button(804, 532, 140, 38, "Back")
    svg.button(956, 532, 148, 38, "Summary", primary=True)
    svg.rect(64, 610, 1072, 58, fill="amberSoft", stroke="amberSoft", rx=20)
    svg.text(92, 644, "Good / Better / Best smart bundle recommendations", cls="h3")


def draw_quote_summary(svg: Svg):
    public_shell(svg, "Get Estimate")
    svg.text(64, 150, "Quotation Summary", cls="h1")
    svg.text(64, 176, "Printable customer quotation view with selected products and services.", cls="body", color="muted")
    svg.rect(64, 210, 510, 116, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(92, 246, "Customer and Vehicle", cls="h2")
    svg.multiline("Customer name, phone number, vehicle model, year, plate number, and service concern.", 92, 272, 420, cls="body", color="muted")
    svg.rect(606, 210, 530, 116, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(636, 246, "Quotation Actions", cls="h2")
    svg.button(636, 270, 110, 38, "Save", primary=True)
    svg.button(760, 270, 120, 38, "Print")
    svg.button(894, 270, 120, 38, "Reset")
    svg.table(64, 356, 690, 230, ["Line item", "Qty", "Unit", "Total"], rows=5)
    svg.rect(784, 356, 352, 230, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(818, 394, "Totals", cls="h2")
    for i, (label, value) in enumerate([("Parts subtotal", "PHP 5,950"), ("Services subtotal", "PHP 2,500"), ("Estimated total", "PHP 8,450")]):
        y = 430 + i * 46
        svg.text(818, y, label, cls="body", color="muted")
        svg.text(1050, y, value, cls="h3", anchor="end", color="blue" if i == 2 else "text")
    svg.button(818, 540, 150, 38, "Open Print Preview", primary=True)


def draw_public_service_orders(svg: Svg):
    public_shell(svg, "Service Orders")
    page_title(svg, 64, 148, "Service Order Workflow", "Handled with structure", "Customer-facing explanation of service order handling.")
    steps = ["Initial Assessment", "Quotation", "Service Handling", "Completion"]
    for i, label in enumerate(steps):
        x = 64 + i * 270
        svg.rect(x, 292, 240, 160, fill="surface", stroke="border", rx=22, shadow=True)
        svg.circle(x + 32, 330, 22, fill="blueSoft" if i % 2 == 0 else "redSoft", stroke="border")
        svg.text(x + 66, 326, f"Step {i + 1}", cls="small", color="muted")
        svg.text(x + 66, 352, label, cls="h3")
        svg.multiline("Short process details shown for customers.", x + 28, 386, 180, cls="body", color="muted")
    svg.rect(64, 492, 520, 104, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(94, 530, "Status Tracking", cls="h2")
    svg.text(94, 556, "Customers can coordinate updates with store staff.", cls="body", color="muted")
    svg.rect(616, 492, 520, 104, fill="darkGrad", stroke="dark", rx=22)
    svg.text(646, 530, "Contact the shop for service coordination", cls="h2", color="surface")
    svg.button(646, 550, 150, 36, "Get Estimate", primary=True)


def draw_public_about(svg: Svg):
    public_shell(svg, "About")
    svg.rect(64, 140, 1072, 210, fill="url(#heroGrad)", stroke="border", rx=28, shadow=True)
    page_title(svg, 96, 184, "About", "Limen Auto Parts Center", "Family-owned auto parts business in Pasay City serving customers with genuine Mitsubishi parts.")
    svg.rect(730, 170, 330, 148, fill="surface2", stroke="border", rx=22)
    svg.text(895, 250, "Shop / Map Visual", cls="h2", anchor="middle", color="muted")
    for i, (label, value) in enumerate([("13 Years", "In service"), ("Pasay City", "Location"), ("2 Floors", "Sales + stockroom")]):
        metric(svg, 64 + i * 366, 380, 330, value, label, fill="surface")
    svg.rect(64, 500, 510, 116, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(94, 536, "Our Story and Operations", cls="h2")
    svg.multiline("Business background, daily operations, sales area, and stockroom process.", 94, 562, 420, cls="body", color="muted")
    svg.rect(606, 500, 530, 116, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(636, 536, "Meet Our Mechanics", cls="h2")
    for i in range(3):
        svg.circle(652 + i * 72, 574, 22, fill="blueSoft", stroke="blueSoft")
        svg.text(652 + i * 72, 604, f"Staff {i + 1}", cls="tiny", anchor="middle", color="muted")


def draw_login(svg: Svg):
    svg.rect(0, 0, WIDTH, HEIGHT, fill="url(#darkGrad)", stroke="dark", rx=0)
    svg.rect(80, 70, 460, 580, fill="surface", stroke="surface", rx=30, shadow=True)
    logo(svg, 126, 116)
    svg.text(126, 214, "Welcome Back", cls="h1")
    svg.text(126, 244, "Sign in to LimenServe MIS", cls="body", color="muted")
    svg.input(126, 292, 320, 44, "Email Address")
    svg.input(126, 356, 320, 44, "Password")
    svg.pill(126, 420, 112, 28, "Remember me", fill="surface2")
    svg.text(322, 439, "Forgot password?", cls="small", color="blue")
    svg.button(126, 476, 320, 44, "Sign In", primary=True)
    svg.rect(126, 546, 320, 62, fill="blueSoft", stroke="blueSoft", rx=16)
    svg.text(150, 572, "Contact Admin", cls="h3", color="blue")
    svg.text(150, 594, "For password concerns, contact administration.", cls="small", color="muted")
    svg.rect(624, 110, 450, 500, fill="surface", stroke="surface", rx=34, opacity="0.10")
    svg.text(700, 280, "LimenServe", cls="h1", color="surface")
    svg.multiline("Internal management workspace for inventory, sales, quotations, service orders, reports, and stockroom operations.", 700, 322, 300, cls="body", color="surface", line_height=18)


def draw_workspace_shell(svg: Svg):
    internal_shell(svg, "Dashboard")
    svg.rect(282, 104, 250, 130, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(310, 142, "Sidebar Navigation", cls="h2")
    svg.text(310, 168, "Role-based module links", cls="body", color="muted")
    svg.rect(558, 104, 250, 130, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(586, 142, "Top Header", cls="h2")
    svg.text(586, 168, "Page title, search, alerts", cls="body", color="muted")
    svg.rect(834, 104, 310, 130, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(862, 142, "Protected Content Area", cls="h2")
    svg.text(862, 168, "Module pages render here", cls="body", color="muted")
    svg.rect(282, 270, 862, 350, fill="surface", stroke="border", rx=24, shadow=True)
    svg.text(320, 316, "Workspace Module Canvas", cls="h1")
    svg.multiline("This shell is shared by the dashboard, inventory, POS, quotation, service orders, stockroom, reports, user management, and CMS screens.", 320, 352, 680, cls="body", color="muted", line_height=18)


def draw_dashboard(svg: Svg):
    internal_shell(svg, "Dashboard")
    svg.rect(282, 100, 862, 110, fill="darkGrad", stroke="dark", rx=24)
    svg.text(316, 146, "Dashboard Overview", cls="h1", color="surface")
    svg.text(316, 174, "Forecast refresh, reports shortcut, and analytics status.", cls="body", color="surface")
    for i, (label, value, fill) in enumerate([
        ("Predicted Revenue", "PHP 84K", "blueSoft"),
        ("Forecasted Units", "128", "greenSoft"),
        ("Top Selling Items", "24", "amberSoft"),
        ("Upsell Opportunities", "17", "purpleSoft"),
    ]):
        metric(svg, 282 + i * 218, 236, 198, label, value, fill=fill)
    svg.rect(282, 346, 520, 240, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(312, 384, "Item Sales Trend", cls="h2")
    svg.line(330, 540, 760, 540, color="border")
    svg.line(330, 410, 330, 540, color="border")
    svg.path("M340 520 C410 486 450 500 510 450 S640 438 760 386", color="blue", width=4)
    svg.rect(828, 346, 316, 240, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(858, 384, "Recent Transactions", cls="h2")
    for i in range(4):
        y = 414 + i * 42
        svg.circle(862, y, 10, fill="blueSoft", stroke="blueSoft")
        svg.text(882, y + 4, f"Sale / service activity {i + 1}", cls="body")
        svg.text(1086, y + 4, "Today", cls="tiny", color="muted", anchor="end")
    svg.rect(828, 608, 316, 58, fill="redSoft", stroke="redSoft", rx=18)
    svg.text(858, 642, "Low-stock alerts shown here", cls="h3", color="red")


def draw_inventory(svg: Svg):
    internal_shell(svg, "Inventory")
    svg.input(282, 104, 370, 42, "Search products, SKU, barcode")
    svg.button(668, 104, 110, 42, "Scan")
    svg.button(792, 104, 120, 42, "Add Stock", primary=True)
    svg.button(926, 104, 130, 42, "Import Price")
    svg.button(1070, 104, 74, 42, "Print")
    svg.table(282, 174, 560, 390, ["SKU", "Part Name", "Stock", "Price", "Status"], rows=7)
    svg.rect(868, 174, 276, 158, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(898, 212, "Filters", cls="h2")
    svg.pill(898, 232, 94, 28, "Category")
    svg.pill(1006, 232, 86, 28, "Stock")
    svg.pill(898, 274, 118, 28, "Compatibility")
    svg.rect(868, 358, 276, 206, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(898, 396, "Product Label Preview", cls="h2")
    svg.rect(898, 420, 216, 78, fill="surface2", stroke="border2", rx=12)
    for i in range(7):
        svg.line(916 + i * 24, 486, 916 + i * 24, 444, color="dark", width=2)
    svg.button(898, 516, 120, 34, "Print Label", primary=True)
    svg.rect(282, 592, 862, 56, fill="blueSoft", stroke="blueSoft", rx=18)
    svg.text(314, 626, "Add stock and product-edit modals open above the inventory table.", cls="h3", color="blue")


def draw_pos(svg: Svg):
    internal_shell(svg, "Point of Sale")
    svg.input(282, 104, 456, 42, "Scan barcode or search product")
    svg.button(754, 104, 110, 42, "New Sale", primary=True)
    svg.button(878, 104, 90, 42, "Hold")
    svg.button(982, 104, 90, 42, "Clear")
    svg.rect(282, 176, 500, 396, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 214, "Product Results", cls="h2")
    for i, label in enumerate(["Oil Filter", "Brake Pad", "Spark Plug", "Cabin Filter"]):
        product_card(svg, 314 + (i % 2) * 220, 244 + (i // 2) * 132, 196, 112, label, f"PHP {750 + i * 450:,}")
    svg.rect(812, 176, 332, 396, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(844, 214, "Current Cart", cls="h2")
    svg.table(844, 238, 268, 186, ["Item", "Qty", "Total"], rows=4)
    svg.text(844, 466, "Payment Method", cls="h3")
    svg.pill(844, 482, 70, 28, "Cash", fill="greenSoft", color="green")
    svg.pill(926, 482, 76, 28, "GCash")
    svg.text(844, 546, "Total: PHP 6,250", cls="h2", color="blue")
    svg.button(964, 520, 148, 42, "Checkout", primary=True)


def draw_quotation(svg: Svg):
    internal_shell(svg, "Quotation")
    svg.rect(282, 104, 408, 150, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 140, "Customer and Vehicle", cls="h2")
    svg.input(314, 160, 160, 34, "Customer")
    svg.input(490, 160, 160, 34, "Phone")
    svg.input(314, 206, 160, 34, "Vehicle")
    svg.input(490, 206, 160, 34, "Year")
    svg.rect(718, 104, 426, 150, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(750, 140, "Quote Actions", cls="h2")
    svg.button(750, 164, 90, 36, "Save", primary=True)
    svg.button(854, 164, 90, 36, "Print")
    svg.button(958, 164, 110, 36, "Lookup")
    svg.rect(282, 282, 408, 280, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 320, "Parts and Services Selector", cls="h2")
    svg.input(314, 342, 300, 34, "Search parts")
    for i in range(3):
        svg.rect(314, 398 + i * 48, 320, 34, fill="surface2", stroke="border", rx=10)
        svg.text(332, 420 + i * 48, f"Selectable quote line {i + 1}", cls="body")
    svg.table(718, 282, 426, 280, ["Description", "Qty", "Amount"], rows=5)
    svg.rect(718, 588, 426, 62, fill="blueSoft", stroke="blueSoft", rx=18)
    svg.text(746, 626, "Computed quotation total: PHP 8,450", cls="h2", color="blue")


def draw_services_admin(svg: Svg):
    internal_shell(svg, "Service Orders")
    for i, label in enumerate(["Pending", "In Progress", "Completed", "Cancelled"]):
        svg.pill(282 + i * 120, 104, 104, 32, label, fill="blueSoft" if i == 0 else "surface")
    svg.button(1014, 102, 130, 36, "New Order", primary=True)
    svg.table(282, 164, 550, 420, ["Customer", "Vehicle", "Concern", "Status"], rows=8)
    svg.rect(858, 164, 286, 420, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(890, 202, "Order Detail", cls="h2")
    svg.multiline("Diagnosis, parts needed, labor notes, assigned mechanic, and progress controls.", 890, 234, 218, cls="body", color="muted")
    for i, label in enumerate(["Assign", "Update Status", "Complete"]):
        svg.button(890, 338 + i * 52, 142, 36, label, primary=i == 1)
    svg.rect(282, 608, 862, 48, fill="greenSoft", stroke="greenSoft", rx=18)
    svg.text(314, 638, "Service modals support create, assign, complete, and update workflows.", cls="h3", color="green")


def draw_stockroom(svg: Svg):
    internal_shell(svg, "3D Stockroom")
    svg.rect(282, 104, 570, 470, fill="dark", stroke="dark", rx=22, shadow=True)
    svg.text(316, 146, "3D Stockroom View", cls="h2", color="surface")
    for i in range(5):
        x = 330 + i * 84
        svg.rect(x, 210, 48, 250, fill="#334155", stroke="#475569", rx=6)
        for j in range(4):
            svg.rect(x + 6, 224 + j * 52, 36, 30, fill="#64748b", stroke="#94a3b8", rx=4)
    svg.path("M358 512 C440 480 520 500 620 430 S725 384 804 330", color="red", width=4)
    svg.circle(804, 330, 12, fill="red", stroke="red")
    svg.rect(878, 104, 266, 100, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 142, "Floor Controls", cls="h2")
    svg.pill(910, 158, 70, 28, "Floor 1", fill="blueSoft", color="blue")
    svg.pill(994, 158, 70, 28, "Floor 2")
    svg.rect(878, 230, 266, 110, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 268, "Product Search", cls="h2")
    svg.input(910, 284, 190, 34, "Find item")
    svg.rect(878, 366, 266, 208, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 404, "Location Details", cls="h2")
    svg.multiline("Shelf B-04, Level 2, Slot 16\nRoute steps and distance.", 910, 432, 190, cls="body", color="muted")
    svg.button(910, 512, 142, 36, "Start Route", primary=True)
    svg.rect(282, 604, 862, 46, fill="surface", stroke="border", rx=18)
    svg.text(314, 633, "Route guidance and stockroom legend", cls="h3")


def draw_parts_mapping(svg: Svg):
    internal_shell(svg, "Parts Mapping")
    svg.rect(282, 104, 568, 470, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 144, "3D Layout Editor Canvas", cls="h2")
    svg.rect(330, 200, 440, 280, fill="surface2", stroke="border2", rx=16)
    for i in range(4):
        svg.rect(370 + i * 86, 244, 50, 170, fill="blueSoft", stroke="blue", rx=8)
    svg.rect(878, 104, 266, 92, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 142, "Editor Tools", cls="h2")
    for i, label in enumerate(["Select", "Move", "Add"]):
        svg.pill(910 + i * 72, 156, 62, 26, label, fill="surface2")
    svg.rect(878, 222, 266, 168, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 260, "Structure Tree", cls="h2")
    for i, label in enumerate(["Floors", "Zones", "Aisles", "Shelves", "Slots"]):
        svg.text(928, 292 + i * 22, label, cls="body")
    svg.rect(878, 416, 266, 158, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(910, 454, "Properties", cls="h2")
    svg.input(910, 474, 190, 32, "Name")
    svg.input(910, 516, 190, 32, "Position")
    svg.rect(282, 604, 862, 46, fill="blueSoft", stroke="blueSoft", rx=18)
    svg.text(314, 633, "Save layout and publish changes to the stockroom locator.", cls="h3", color="blue")


def draw_reports(svg: Svg):
    internal_shell(svg, "Reports")
    svg.pill(282, 104, 110, 32, "Sales Report", fill="blueSoft", color="blue")
    svg.pill(406, 104, 136, 32, "Inventory Report")
    svg.input(718, 100, 160, 38, "Date range")
    svg.button(892, 100, 104, 38, "Export")
    svg.button(1010, 100, 134, 38, "Encode Sale", primary=True)
    for i, (label, value) in enumerate([("Total Sales", "PHP 52K"), ("Units Sold", "87"), ("Low Stock", "15"), ("Service Revenue", "PHP 9K")]):
        metric(svg, 282 + i * 218, 164, 198, label, value)
    svg.rect(282, 274, 520, 250, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 312, "Sales Trend Chart", cls="h2")
    svg.path("M330 474 C404 428 448 452 520 392 S650 376 760 326", color="blue", width=4)
    svg.rect(828, 274, 316, 250, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(858, 312, "Report Summary", cls="h2")
    for i in range(4):
        svg.rect(858, 342 + i * 42, 236, 28, fill="surface2", stroke="border", rx=8)
    svg.table(282, 550, 862, 112, ["Date", "Product", "Qty", "Amount", "Encoded By"], rows=2)


def draw_users(svg: Svg):
    internal_shell(svg, "User Management")
    for i, (label, value, fill) in enumerate([("Total Users", "12", "blueSoft"), ("Active Staff", "10", "greenSoft"), ("Administrators", "2", "amberSoft")]):
        metric(svg, 282 + i * 290, 104, 260, label, value, fill=fill)
    svg.input(282, 214, 440, 42, "Search users by name, email, or role")
    svg.button(1014, 214, 130, 42, "Add User", primary=True)
    svg.table(282, 286, 550, 306, ["User", "Email", "Role", "Status"], rows=6)
    svg.rect(858, 286, 286, 306, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(890, 324, "Add / Edit User", cls="h2")
    svg.input(890, 350, 210, 34, "Full name")
    svg.input(890, 394, 210, 34, "Email")
    svg.input(890, 438, 210, 34, "Role")
    svg.input(890, 482, 210, 34, "Password")
    svg.button(890, 538, 124, 36, "Save User", primary=True)


def draw_cms(svg: Svg):
    internal_shell(svg, "Content CMS")
    svg.rect(282, 104, 240, 472, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(314, 142, "Page List", cls="h2")
    for i, label in enumerate(["Home", "About", "Catalog", "Service Orders"]):
        svg.rect(314, 174 + i * 46, 168, 34, fill="blueSoft" if i == 0 else "surface2", stroke="border", rx=10)
        svg.text(334, 196 + i * 46, label, cls="body", color="blue" if i == 0 else "text")
    svg.rect(550, 104, 332, 472, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(582, 142, "Section Editor", cls="h2")
    svg.input(582, 174, 252, 34, "Hero title")
    svg.input(582, 222, 252, 34, "Hero subtitle")
    svg.rect(582, 274, 252, 142, fill="surface2", stroke="border", rx=14)
    svg.text(606, 304, "Card / FAQ / contact blocks", cls="h3")
    svg.button(582, 442, 120, 36, "Save Draft")
    svg.button(714, 442, 120, 36, "Publish", primary=True)
    svg.rect(910, 104, 234, 472, fill="surface", stroke="border", rx=22, shadow=True)
    svg.text(942, 142, "Live Preview", cls="h2")
    svg.rect(942, 174, 170, 330, fill="url(#heroGrad)", stroke="border", rx=18)
    svg.text(970, 230, "Public page", cls="h3")
    svg.text(970, 254, "preview", cls="body", color="muted")
    svg.rect(282, 608, 862, 46, fill="blueSoft", stroke="blueSoft", rx=18)
    svg.text(314, 637, "Manage public content sections, preview changes, and publish updates.", cls="h3", color="blue")


DRAWERS = {
    "public_home": draw_public_home,
    "public_catalog": draw_public_catalog,
    "estimate_choice": draw_estimate_choice,
    "estimate_builder": draw_estimate_builder,
    "quote_summary": draw_quote_summary,
    "public_service_orders": draw_public_service_orders,
    "public_about": draw_public_about,
    "login": draw_login,
    "workspace_shell": draw_workspace_shell,
    "dashboard": draw_dashboard,
    "inventory": draw_inventory,
    "pos": draw_pos,
    "quotation": draw_quotation,
    "services_admin": draw_services_admin,
    "stockroom": draw_stockroom,
    "parts_mapping": draw_parts_mapping,
    "reports": draw_reports,
    "users": draw_users,
    "cms": draw_cms,
}


def generate_svgs():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    manifest = []
    for index, fig in enumerate(FIGURES, start=1):
        number = f"{index:02d}"
        filename = f"figure-{number}-{fig['slug']}"
        svg = Svg()
        DRAWERS[fig["layout"]](svg)
        svg_path = FIG_DIR / f"{filename}.svg"
        svg_path.write_text(svg.finish(), encoding="utf-8")
        manifest.append(
            {
                "number": index,
                "title": fig["title"],
                "description": fig["description"].format(n=index),
                "svg": str(svg_path),
                "png": str(FIG_DIR / f"{filename}.png"),
            }
        )
    (OUT_DIR / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    lines = ["# LimenServe Mid-Fidelity System Wireframe Figure Descriptions", ""]
    for item in manifest:
        lines.append(f"## Figure {item['number']}. {item['title']}")
        lines.append(item["description"])
        lines.append("")
    (OUT_DIR / "figure-descriptions.md").write_text("\n".join(lines), encoding="utf-8")
    return manifest


def build_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    manifest_path = OUT_DIR / "manifest.json"
    if not manifest_path.exists():
        raise SystemExit("Run the svg command first.")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    missing = [item["png"] for item in manifest if not Path(item["png"]).exists()]
    if missing:
        raise SystemExit("Missing rendered PNG files:\n" + "\n".join(missing))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LIMENSERVE SYSTEM MID-FIDELITY WIREFRAMES")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "The following colored mid-fidelity wireframes present the main public and internal interfaces of LimenServe. "
        "The figures remove the browser address bar and focus on the application screens, using the system's white, gray, blue, and red visual style."
    )

    for item in manifest:
        doc.add_paragraph()
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(item["png"], width=Inches(6.85))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cap.add_run(f"Figure {item['number']}. ")
        r1.bold = True
        r2 = cap.add_run(item["title"])
        r1.font.name = r2.font.name = "Times New Roman"
        r1.font.size = r2.font.size = Pt(11)

        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc.paragraph_format.first_line_indent = Inches(0.25)
        desc.add_run(item["description"])
        if item["number"] != manifest[-1]["number"]:
            doc.add_page_break()

    doc_path = OUT_DIR / "LimenServe_System_Wireframes_Mid_Fidelity.docx"
    doc.save(doc_path)
    return doc_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["svg", "docx", "all"])
    args = parser.parse_args()
    if args.command in {"svg", "all"}:
        generate_svgs()
    if args.command in {"docx", "all"}:
        print(build_docx())


if __name__ == "__main__":
    main()
