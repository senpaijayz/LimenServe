from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Appendix_Gantt.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Appendix_Gantt_ContextDiagram.docx"
DIAGRAM_IMAGE = ROOT / "Docs" / "limenserve_context_diagram.png"


def insert_paragraph_before(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else None,
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text: str, fnt, max_width: int):
    lines = []
    for part in text.split("\n"):
        words = part.split()
        current = ""
        for word in words:
            probe = f"{current} {word}".strip()
            if draw.textbbox((0, 0), probe, font=fnt)[2] <= max_width or not current:
                current = probe
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def draw_centered(draw, box, text, fnt, fill=(25, 25, 25), spacing=6):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, int(x2 - x1 - 26))
    heights = [
        draw.textbbox((0, 0), line, font=fnt)[3] - draw.textbbox((0, 0), line, font=fnt)[1]
        for line in lines
    ]
    total_h = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + spacing


def draw_box(draw, box, title, body=""):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, outline=(20, 20, 20), width=4, fill=(248, 248, 248))
    if body:
        draw_centered(draw, (x1 + 12, y1 + 18, x2 - 12, y1 + 78), title, font(27, True))
        draw.line((x1 + 24, y1 + 86, x2 - 24, y1 + 86), fill=(80, 80, 80), width=2)
        draw_centered(draw, (x1 + 20, y1 + 94, x2 - 20, y2 - 14), body, font(21))
    else:
        draw_centered(draw, (x1 + 12, y1 + 16, x2 - 12, y2 - 16), title, font(26, True))


def arrow(draw, start, end, label, label_offset=(0, 0)):
    draw.line((start, end), fill=(45, 45, 45), width=4)
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    p1 = (ex, ey)
    p2 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
    p3 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=(45, 45, 45))

    mid = ((sx + ex) / 2 + label_offset[0], (sy + ey) / 2 + label_offset[1])
    fnt = font(19)
    lines = wrap_text(draw, label, fnt, 260)
    line_h = 22
    pad = 8
    max_w = max(draw.textbbox((0, 0), line, font=fnt)[2] for line in lines)
    box = (
        mid[0] - max_w / 2 - pad,
        mid[1] - (line_h * len(lines)) / 2 - pad,
        mid[0] + max_w / 2 + pad,
        mid[1] + (line_h * len(lines)) / 2 + pad,
    )
    draw.rounded_rectangle(box, radius=8, fill=(255, 255, 255), outline=(170, 170, 170), width=1)
    y = box[1] + pad
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((mid[0] - w / 2, y), line, font=fnt, fill=(25, 25, 25))
        y += line_h


def generate_context_diagram():
    img = Image.new("RGB", (1700, 2200), "white")
    d = ImageDraw.Draw(img)
    black = (25, 25, 25)

    title = "Appendix 4. Level 1: Context Diagram of LimenServe"
    d.text((130, 80), title, font=font(34, True), fill=black)
    subtitle = "LimenServe External Entities and Data Flow"
    w = d.textbbox((0, 0), subtitle, font=font(30, True))[2]
    d.text(((1700 - w) / 2, 175), subtitle, font=font(30, True), fill=black)

    center = (625, 870, 1075, 1160)
    admin = (115, 435, 505, 650)
    customer = (1195, 435, 1585, 650)
    staff = (115, 1390, 505, 1605)
    supplier = (1195, 1390, 1585, 1605)

    draw_box(
        d,
        center,
        "1.0\nLimenServe System",
        "Inventory, quotations, POS transactions, service orders, reports, user accounts, and stockroom visualization",
    )
    draw_box(d, admin, "Owner / Administrator", "User accounts, product records, reports, and system monitoring")
    draw_box(d, customer, "Customer", "Product inquiry, quotation request, purchase details, and service order request")
    draw_box(d, staff, "Cashier / Staff", "Sales transaction, quotation encoding, service update, and inventory checking")
    draw_box(d, supplier, "Supplier", "Product delivery details, item information, and restocking support")

    # Bidirectional data flows with clear offsets to avoid overlap.
    arrow(d, (505, 500), (625, 900), "Product data, account control, report request", (-96, -30))
    arrow(d, (625, 965), (505, 585), "Reports, dashboard, inventory status", (104, 18))

    arrow(d, (1195, 500), (1075, 900), "Inquiry, order, quotation request", (100, -30))
    arrow(d, (1075, 965), (1195, 585), "Quotation, receipt, service status", (-105, 18))

    arrow(d, (505, 1510), (625, 1115), "Sales, service order, inventory update", (-96, 32))
    arrow(d, (625, 1050), (505, 1435), "Receipt, stock status, transaction record", (105, -20))

    arrow(d, (1195, 1510), (1075, 1115), "Delivery details and product info", (98, 32))
    arrow(d, (1075, 1050), (1195, 1435), "Restock request and item requirements", (-105, -20))

    note = (
        "The context diagram shows the overall boundary of LimenServe and the flow of information "
        "between the system and its external users."
    )
    d.rounded_rectangle((235, 1785, 1465, 1885), radius=18, outline=(120, 120, 120), width=2, fill=(250, 250, 250))
    draw_centered(d, (260, 1795, 1440, 1875), note, font(23), fill=black)

    caption = "Figure A2. Level 1 Context Diagram of LimenServe"
    cw = d.textbbox((0, 0), caption, font=font(28))[2]
    d.text(((1700 - cw) / 2, 1950), caption, font=font(28), fill=black)

    img.save(DIAGRAM_IMAGE, quality=95)


def add_context_appendix(doc: Document):
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Appendix A. Interview Report":
            anchor = paragraph
            break
    if anchor is None:
        raise SystemExit("Could not find Appendix A insertion point.")

    generate_context_diagram()
    paragraph = insert_paragraph_before(anchor, "")
    paragraph.paragraph_format.page_break_before = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(DIAGRAM_IMAGE), width=Inches(6.55))


def main():
    doc = Document(SOURCE_DOCX)
    if any("Level 1: Context Diagram of LimenServe" in p.text for p in doc.paragraphs):
        raise SystemExit("Context diagram appendix already exists.")
    add_context_appendix(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
