from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Appendix_Gantt_ContextDiagram.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Editable_Gantt_ContextDiagram.docx"

RESEARCHERS = (
    "Jay Iverson A. Dela Cruz, Alden M. Suacillo, Don Rich M. Ulanday, "
    "Christian Dave D. Vargas, Estefhano C. Villafuerte"
)
PROGRAM = "Bachelor of Science in Information Technology"
TITLE = (
    "LimenServe: A Web-Based Management Information System with Automated Cost "
    "Estimation and Quotation for Services and Orders of Limen Auto Parts Center"
)

MONTHS = [
    ("2025", "Oct"),
    ("2025", "Nov"),
    ("2025", "Dec"),
    ("2026", "Jan"),
    ("2026", "Feb"),
    ("2026", "Mar"),
    ("2026", "Apr"),
    ("2026", "May"),
    ("2026", "Jun"),
    ("2026", "Jul"),
    ("2026", "Aug"),
    ("2026", "Sep"),
    ("2026", "Oct"),
    ("2026", "Nov"),
    ("2026", "Dec"),
    ("2027", "Jan"),
    ("2027", "Feb"),
]

SCHEDULE = [
    ("Proposal drafting", 0, 2),
    ("Proposal defense", 3, 3),
    ("System design", 4, 5),
    ("System development", 5, 14),
    ("System testing", 12, 14),
    ("System evaluation", 13, 14),
    ("System implementation", 13, 14),
    ("Manuscript preparation", 0, 15),
    ("Final oral defense", 15, 15),
    ("Manuscript review", 14, 16),
]


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def insert_paragraph_before(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_before(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addprevious(table._tbl)
    return table


def remove_paragraph(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=2):
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=45, start=45, bottom=45, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_width(cell, inches: float):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[float]):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl.insert(0, grid)


def set_cell_text(
    cell,
    text: str,
    *,
    size=7,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    fill: str | None = None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def set_no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tc_pr.append(borders)


def add_info_table(anchor):
    table = insert_table_before(anchor, 3, 3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name of Researcher(s)", ":", RESEARCHERS),
        ("Program", ":", PROGRAM),
        ("Title of Study", ":", TITLE),
    ]
    widths = [1.45, 0.15, 4.75]
    set_table_grid(table, widths)
    for row, values in zip(table.rows, rows):
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_text(
                cell,
                value,
                size=7.5,
                italic=(idx == 0),
                align=WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_width(cell, widths[idx])
            set_no_borders(cell)
    return table


def add_gantt_table(anchor):
    rows = 1 + len(SCHEDULE)
    cols = 1 + len(MONTHS)
    table = insert_table_before(anchor, rows, cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [1.55] + [0.29] * len(MONTHS)
    set_table_grid(table, widths)

    header_fill = "E7E7E7"
    bar_fill = "6F8F72"

    set_cell_text(table.cell(0, 0), "ACTIVITY", size=7, bold=True, fill=header_fill)
    set_width(table.cell(0, 0), widths[0])
    for idx, (year, month) in enumerate(MONTHS, start=1):
        header = f"{month}\n{year}"
        set_cell_text(table.cell(0, idx), header, size=4.8, bold=True, fill=header_fill)
        set_width(table.cell(0, idx), widths[idx])

    for row_idx, (activity, start, end) in enumerate(SCHEDULE, start=1):
        set_cell_text(
            table.cell(row_idx, 0),
            activity,
            size=5.7,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_width(table.cell(row_idx, 0), widths[0])
        for month_idx in range(len(MONTHS)):
            fill = bar_fill if start <= month_idx <= end else None
            set_cell_text(table.cell(row_idx, month_idx + 1), "", size=5, fill=fill)
            set_width(table.cell(row_idx, month_idx + 1), widths[month_idx + 1])

    return table


def find_gantt_image_paragraph(doc: Document):
    after_appendix_2 = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Appendix 2. Sprint Backlog":
            after_appendix_2 = True
            continue
        if after_appendix_2 and has_drawing(paragraph):
            return paragraph
        if after_appendix_2 and text == "Appendix A. Interview Report":
            break
    raise SystemExit("Could not locate Gantt image paragraph.")


def main():
    doc = Document(SOURCE_DOCX)
    image_paragraph = find_gantt_image_paragraph(doc)

    title = insert_paragraph_before(image_paragraph, "Appendix 3. Gantt Chart")
    title.paragraph_format.page_break_before = True
    format_paragraph(title, size=10, bold=True, after=8)

    for text, size, bold, after in [
        ("Republic of the Philippines", 8, False, 1),
        ("CAVITE STATE UNIVERSITY", 13, True, 1),
        ("Bacoor City Campus", 9, True, 1),
        ("SHIV, Molino VI, City of Bacoor", 8, False, 10),
        ("DEPARTMENT OF COMPUTER STUDIES", 10, True, 8),
        ("GANTT CHART", 11, True, 8),
    ]:
        p = insert_paragraph_before(image_paragraph, text)
        format_paragraph(p, size=size, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER, after=after)

    add_info_table(image_paragraph)
    spacer = insert_paragraph_before(image_paragraph, "")
    spacer.paragraph_format.space_after = Pt(6)
    add_gantt_table(image_paragraph)
    caption = insert_paragraph_before(image_paragraph, "Figure A1. Gantt Chart of LimenServe development schedule")
    format_paragraph(caption, size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    remove_paragraph(image_paragraph)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
