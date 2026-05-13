from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_Figure_Descriptions.docx"


DESCRIPTION_FIXES = {
    "Figure 6 shows the public service-order information page.": "Figure 3.11 shows the public service-order information page.",
    "Figure 7 shows the public about page of Limen Auto Parts Center.": "Figure 3.12 shows the public about page of Limen Auto Parts Center.",
    "Figure 8 shows the staff login page.": "Figure 3.13 shows the staff login page.",
    "Figure 9 shows the common internal workspace shell used after authentication.": "Figure 3.14 shows the common internal workspace shell used after authentication.",
    "Figure 10 shows the dashboard module.": "Figure 3.15 shows the dashboard module.",
    "Figure 11 shows the inventory management module.": "Figure 3.16 shows the inventory management module.",
    "Figure 12 shows the point-of-sale module.": "Figure 3.17 shows the point-of-sale module.",
    "Figure 13 shows the internal quotation builder used by staff.": "Figure 3.18 shows the internal quotation builder used by staff.",
    "Figure 14 shows the internal service-order management module.": "Figure 3.19 shows the internal service-order management module.",
    "Figure 15 shows the stockroom visualization module.": "Figure 3.20 shows the stockroom visualization module.",
    "Figure 16 shows the parts mapping administration interface.": "Figure 3.21 shows the parts mapping administration interface.",
    "Figure 17 shows the reporting and analytics module.": "Figure 3.22 shows the reporting and analytics module.",
    "Figure 18 shows the user management module.": "Figure 3.23 shows the user management module.",
    "Figure 19 shows the content management module.": "Figure 3.24 shows the content management module.",
}


DESCRIPTIONS = {
    "Figure 1. Conceptual framework of the study": (
        "Figure 1 shows the conceptual framework of the study using the input-process-output model. "
        "It explains how the identified business problems, system requirements, development activities, "
        "and evaluation procedures lead to the proposed LimenServe system output."
    ),
    "Figure 3.1. HIPO diagram of the LimenServe": (
        "Figure 3.1 shows the hierarchical input-process-output structure of LimenServe. "
        "It organizes the major modules of the system and presents how each function supports the overall operation of the platform."
    ),
    "Figure 3.2. Level 0: Context diagram of LimenServe": (
        "Figure 3.2 shows the Level 0 context diagram of LimenServe. It identifies the main external entities, "
        "such as the customer, admin or owner, sales staff, and stock personnel, and shows how information flows between them and the system."
    ),
    "Figure 3.3. System architecture of LimenServe": (
        "Figure 3.3 shows the system architecture of LimenServe. It presents the relationship between the presentation layer, "
        "application layer, and data layer, showing how users interact with system functions and how data is stored and managed."
    ),
    "Figure 3.4. ERD of LimenServe (Crow’s Foot)": (
        "Figure 3.4 shows the entity relationship diagram of LimenServe using Crow's Foot notation. "
        "It presents the database entities, key fields, and relationships used to manage users, customers, products, quotations, sales, service orders, inventory, and stockroom locations."
    ),
    "Figure 3.5. UML use case diagram of LimenServe": (
        "Figure 3.5 shows the UML use case diagram of LimenServe. It presents the main actors and the system functions they can perform, "
        "including browsing products, requesting quotations, managing users, processing sales, managing inventory, viewing reports, and recording service orders."
    ),
    "Figure 3.25. Current technical situation of Limen Auto Parts Center": (
        "Figure 3.25 shows the current technical situation of Limen Auto Parts Center. "
        "It illustrates the manual flow of customer inquiries, price checking, physical stock verification, handwritten recording, and delayed report preparation."
    ),
    "Figure 3.26. Fishbone diagram for problem 1": (
        "Figure 3.26 shows the fishbone analysis for poor credential and data management. "
        "It identifies causes related to materials, methods, machinery, and manpower, including manual consolidation, physical records, and the absence of a centralized database."
    ),
    "Figure 3.27. Fishbone diagram for problem 2": (
        "Figure 3.27 shows the fishbone analysis for inefficient record monitoring. "
        "It explains the causes of scattered records, separate branch storage, isolated files, and manual tracking that make monitoring less organized."
    ),
    "Figure 3.28. Fishbone diagram for problem 3": (
        "Figure 3.28 shows the fishbone analysis for manual enrollment processing. "
        "It presents the causes connected to paper-based forms, manual document verification, and the lack of an automated processing system."
    ),
    "Figure A1. Gantt Chart of LimenServe development schedule": (
        "Figure A1 shows the Gantt chart for the development schedule of LimenServe. "
        "It presents the planned timeline for proposal preparation, system design, development, testing, evaluation, implementation, manuscript preparation, defense, and review."
    ),
    "Figure A2. Level 1 Context Diagram of Limen Serve": (
        "Figure A2 shows the Level 1 context diagram of LimenServe. "
        "It provides a more detailed view of the external entities and the data exchanged between the system, the owner or administrator, customers, staff, and suppliers."
    ),
}


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para.add_run(text)
    return new_para


def format_description(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def next_nonempty(doc: Document, index: int):
    for paragraph in doc.paragraphs[index + 1 :]:
        if paragraph.text.strip():
            return paragraph
    return None


def main():
    doc = Document(SOURCE_DOCX)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for old, new in DESCRIPTION_FIXES.items():
            if text.startswith(old):
                paragraph.text = paragraph.text.replace(old, new, 1)
                format_description(paragraph)
                break

    # Insert from bottom to top so paragraph indexes and insertion points remain stable.
    insertions = []
    for index, paragraph in enumerate(doc.paragraphs):
        caption = paragraph.text.strip()
        if caption in DESCRIPTIONS:
            following = next_nonempty(doc, index)
            expected_start = DESCRIPTIONS[caption].split(". ")[0]
            if following is not None and following.text.strip().startswith(expected_start):
                continue
            insertions.append((index, caption))

    for index, caption in reversed(insertions):
        paragraph = doc.paragraphs[index]
        inserted = insert_paragraph_after(paragraph, DESCRIPTIONS[caption])
        format_description(inserted)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(f"Inserted descriptions: {len(insertions)}")


if __name__ == "__main__":
    main()
