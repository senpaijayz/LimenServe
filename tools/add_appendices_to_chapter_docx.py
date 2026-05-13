from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_With_References.docx"
OUT_DOCX = ROOT / "Docs" / "LIMENSERVECHAPTER1TO3_With_References_Appendices.docx"


PRODUCT_BACKLOG = [
    ("User Authentication", "Admin and staff log in securely based on assigned account credentials.", "All", "High"),
    ("User Management", "Admin manages user accounts, roles, and access permissions.", "Admin", "High"),
    ("Public Catalog", "Customers browse available auto parts and services through the public page.", "Customer", "Medium/High"),
    ("Automated Cost Estimation", "Customers request an estimate by selecting parts, services, and vehicle details.", "Customer", "High"),
    ("Quotation Management", "Admin or staff reviews customer requests and generates quotation records.", "Admin/Staff", "High"),
    ("Inventory Management", "Admin or stock personnel manage product details, prices, quantities, and stock status.", "Admin/Stock Personnel", "High"),
    ("Barcode Product Identification", "Staff scans barcode labels to identify products and update stock records faster.", "Admin/Stock Personnel", "Medium/High"),
    ("Point of Sale", "Cashier processes sales transactions, computes totals, and prepares receipt records.", "Cashier", "High"),
    ("Service Order Management", "Staff records, updates, and monitors service-related customer requests.", "Admin/Service Staff", "High"),
    ("Reports and Analytics", "Admin views sales reports, inventory reports, quotation summaries, and dashboard statistics.", "Admin", "High"),
    ("3D Stockroom Visualization", "Staff uses the 3D stockroom view to locate products in the storage area.", "Admin/Stock Personnel", "Medium"),
    ("Content Management", "Admin updates public page content, service information, and shop details.", "Admin", "Medium"),
    ("System Testing", "Researchers test system modules through unit, integration, functional, usability, and acceptance testing.", "Researchers", "High"),
]


SPRINT_BACKLOG = [
    ("Sprint 1", "Establish core system structure, database design, authentication, and role-based access.", "Done", "High"),
    ("Sprint 2", "Develop public product browsing, shop information pages, and customer estimate request flow.", "Done", "Medium/High"),
    ("Sprint 3", "Implement inventory management, product records, stock status, and barcode-assisted identification.", "Done", "High"),
    ("Sprint 4", "Develop automated quotation, parts and service selection, and quotation summary generation.", "Done", "High"),
    ("Sprint 5", "Implement POS transaction processing, receipt support, and inventory deduction after sales.", "Done", "High"),
    ("Sprint 6", "Develop service order management, workflow status updates, and staff monitoring tools.", "In Progress", "High"),
    ("Sprint 7", "Generate reports, dashboards, analytics summaries, and management views.", "Done", "Medium/High"),
    ("Sprint 8", "Implement 3D stockroom visualization, content management, final testing, and deployment preparation.", "In Progress", "Medium/High"),
]


INTERVIEW_QA = [
    (
        "What is the current problem encountered in the shop operations?",
        "The business still relies on manual checking, handwritten records, and staff familiarity when handling product availability, quotations, receipts, and stockroom item location. These practices may cause delays, pricing inconsistencies, and difficulty in preparing accurate reports.",
    ),
    (
        "How often does this problem occur during daily operations?",
        "The problem may occur during regular business transactions, especially when customers ask for available auto parts, prices, estimates, or service-related details. Since staff members still need to check records or physically verify stock, delays can happen repeatedly during busy periods.",
    ),
    (
        "How does the current process affect daily workflow?",
        "The current process can slow down customer service because employees need to move between the sales area and stockroom or search through manual records before confirming item availability. It can also make sales and inventory reports harder to consolidate at the end of the day or week.",
    ),
    (
        "What current systems or processes are used?",
        "The business mainly uses manual records such as receipts, logs, quotations, and handwritten documents. A computer is available for some digital price checking, but several important operations such as stock monitoring, transaction records, service order updates, and report preparation are still not centralized.",
    ),
    (
        "Are there existing solutions that attempt to address the problem?",
        "The current approach depends on staff experience, manual checking, and basic record keeping. While these methods help the shop continue operating, they are limited because records can be scattered, updates may not be immediate, and item location depends heavily on staff familiarity.",
    ),
    (
        "Are there inconsistencies in how staff handle the same task?",
        "Yes. Since the process is not fully systematized, different staff members may check item availability, prepare quotations, or record transactions differently. This can affect consistency, especially when several customers are being assisted at the same time.",
    ),
    (
        "What common concerns are raised by customers?",
        "Customers commonly need faster answers about product availability, estimated cost, service details, and transaction status. Delays may occur when staff still need to confirm prices, check the stockroom, or prepare quotation details manually.",
    ),
    (
        "What mistakes or errors may happen in the current operation?",
        "Possible errors include inaccurate stock counts, delayed stock updates, inconsistent quotation amounts, misplaced item information, repeated manual encoding, and difficulty locating products in the stockroom.",
    ),
    (
        "If a system were developed, what features would be most helpful?",
        "The most helpful features would include inventory monitoring, automated cost estimation and quotation generation, barcode-assisted product identification, POS transaction processing, service order management, reporting and analytics, role-based access, public catalog browsing, and 3D stockroom visualization.",
    ),
    (
        "What suggestions can improve the current process?",
        "A centralized web-based management information system is suggested to connect inventory, quotations, sales, service orders, reports, and stockroom item location in one platform. This would help reduce manual work, improve accuracy, and support faster customer service.",
    ),
]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_paragraph(paragraph, *, size: int = 10, bold: bool = False, alignment=None, first_indent=False):
    if alignment is not None:
        paragraph.alignment = alignment
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def add_page_break(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, *, level: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if level == 1 else 11)
    run.bold = True
    paragraph.paragraph_format.space_after = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    format_paragraph(paragraph, first_indent=True)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
        table.rows[0].cells[idx].width = Inches(widths[idx])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value, size=9)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_product_backlog(doc: Document):
    add_page_break(doc)
    add_heading(doc, "Appendix 1. Product Backlog")
    add_table(
        doc,
        ["FEATURES", "USER STORY/TASK", "ROLES", "PRIORITY"],
        PRODUCT_BACKLOG,
        [1.55, 4.2, 1.45, 1.0],
    )


def add_sprint_backlog(doc: Document):
    add_page_break(doc)
    add_heading(doc, "Appendix 2. Sprint Backlog")
    add_table(
        doc,
        ["SPRINT", "SPRINT GOAL", "STATUS", "PRIORITY"],
        SPRINT_BACKLOG,
        [1.1, 4.8, 1.2, 1.1],
    )


def add_interview_report(doc: Document):
    add_page_break(doc)
    add_heading(doc, "Appendix A. Interview Report")
    add_heading(doc, "INTERVIEW REPORT", level=2)
    add_body(
        doc,
        "Interview with the representative of Limen Auto Parts Center regarding the current sales, inventory, quotation, service order, and stockroom management process.",
    )

    add_heading(doc, "INTERVIEWEE PROFILE", level=2)
    profile = [
        ("Interviewee Name:", "Representative of Limen Auto Parts Center"),
        ("Company:", "Limen Auto Parts Center"),
        ("Company Address:", "Efipanio de los Santos Avenue, Pasay City, Metro Manila"),
        ("Work Position:", "Owner/Authorized Shop Representative"),
        ("Date of Interview:", "2026"),
    ]
    for label, value in profile:
        paragraph = doc.add_paragraph()
        label_run = paragraph.add_run(label + " ")
        label_run.bold = True
        value_run = paragraph.add_run(value)
        for run in (label_run, value_run):
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        paragraph.paragraph_format.space_after = Pt(3)

    add_heading(doc, "INTRODUCTION", level=2)
    add_body(
        doc,
        "This interview report presents the information gathered from Limen Auto Parts Center for the development of LimenServe. The purpose of the interview was to understand the current business process, identify operational problems, and determine the system features needed to support the shop's daily activities.",
    )
    add_body(
        doc,
        "The interview focused on manual inventory checking, quotation preparation, sales recording, service order monitoring, reporting, customer inquiry handling, and stockroom item location. The responses guided the researchers in defining the product backlog, system modules, diagrams, and development priorities of the proposed web-based management information system.",
    )

    add_heading(doc, "INTERVIEW QUESTIONS AND ANSWERS", level=2)
    for index, (question, answer) in enumerate(INTERVIEW_QA, 1):
        q = doc.add_paragraph()
        q_run = q.add_run(f"{index}. {question}")
        q_run.bold = True
        q_run.font.name = "Times New Roman"
        q_run.font.size = Pt(10)
        q.paragraph_format.space_after = Pt(2)
        q.paragraph_format.keep_with_next = True

        a = doc.add_paragraph()
        a_run = a.add_run("Answer: ")
        a_run.bold = True
        a_run.font.name = "Times New Roman"
        a_run.font.size = Pt(10)
        run = a.add_run(answer)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        a.paragraph_format.line_spacing = 1.15
        a.paragraph_format.space_after = Pt(5)

    add_heading(doc, "KEY FINDINGS AND DISCUSSION", level=2)
    add_body(
        doc,
        "Based on the interview, the researchers found that Limen Auto Parts Center needs a more organized way to manage sales, inventory, quotations, service orders, and reports. The current process depends on manual records and staff familiarity, which may affect transaction speed and accuracy. A centralized system can improve the flow of information by allowing authorized users to access updated product, customer, quotation, and transaction records in one platform.",
    )

    add_heading(doc, "CONCLUSION", level=2)
    add_body(
        doc,
        "The interview confirmed the need for LimenServe as a web-based management information system with automated cost estimation and quotation. By integrating inventory monitoring, barcode-assisted product identification, POS transactions, service order management, reports, public browsing, role-based access, and stockroom visualization, the proposed system can help reduce manual work and improve operational control.",
    )

    add_heading(doc, "REFLECTION", level=2)
    add_body(
        doc,
        "As researchers, the group gained a clearer understanding of how manual retail and service operations affect customer service, inventory accuracy, and record management. The interview helped connect the actual needs of Limen Auto Parts Center with the system features that were included in the proposed solution.",
    )
    add_body(
        doc,
        "The information gathered from the interview guided the development of LimenServe and supported the creation of the system diagrams, product backlog, sprint backlog, and system requirements. Overall, the interview helped the researchers design a system that responds to actual operational problems rather than only theoretical requirements.",
    )
    doc.add_paragraph()
    add_body(doc, "Approved by:")
    doc.add_paragraph("\n\n")


def add_placeholder_appendix(doc: Document, title: str):
    add_page_break(doc)
    add_heading(doc, title)
    doc.add_paragraph("\n\n\n\n\n\n")


def main():
    doc = Document(SOURCE_DOCX)
    if any("Appendix 1. Product Backlog" in p.text for p in doc.paragraphs):
        raise SystemExit("The document already appears to contain appendices.")

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_product_backlog(doc)
    add_sprint_backlog(doc)
    add_interview_report(doc)
    add_placeholder_appendix(doc, "Appendix B. Interview Report")
    add_placeholder_appendix(doc, "Appendix C. Concept Checklist")
    add_placeholder_appendix(doc, "Appendix D. Request for Technical Adviser and Tech Critic")
    add_placeholder_appendix(doc, "Appendix E. Approval Sheet")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
